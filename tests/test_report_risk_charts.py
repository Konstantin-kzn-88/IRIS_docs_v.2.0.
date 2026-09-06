import json
from pathlib import Path

import pytest
from docx import Document

from iris_v2.report_risk_charts import (
    FN_EMPTY_TEXT,
    FN_MARKER,
    ReportRiskChartsError,
    prepare_risk_charts,
    render_risk_chart,
)
from iris_v2.risk_summary import build_fg_points, build_fn_points


def write_json(path: Path, value: object) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")


def risk_row(code: str, fatalities: int, frequency: float, damage: float) -> dict:
    return {
        "scenario_code": code,
        "fatalities_count": fatalities,
        "scenario_frequency": frequency,
        "total_damage": damage,
    }


def write_results(tmp_path: Path, rows: list[dict]) -> None:
    write_json(tmp_path / "risk_results.json", {"results": rows})
    write_json(
        tmp_path / "risk_summary.json",
        {
            "case_count": len(rows),
            "risk_unit": "1/год",
            "fg_damage_unit": "млн руб.",
            "fn_points": build_fn_points(rows),
            "fg_points": build_fg_points(rows),
        },
    )


def test_charts_are_regenerated_from_current_points(tmp_path: Path) -> None:
    pytest.importorskip("matplotlib")
    write_results(
        tmp_path,
        [
            risk_row("С1", 1, 1.0e-4, 1000.0),
            risk_row("С2", 2, 5.0e-5, 2500.0),
        ],
    )

    charts = prepare_risk_charts(tmp_path)

    assert charts.fn_path is not None
    assert charts.fg_path is not None
    assert charts.fn_path.read_bytes().startswith(b"\x89PNG\r\n\x1a\n")
    assert charts.fg_path.read_bytes().startswith(b"\x89PNG\r\n\x1a\n")


def test_no_fatalities_uses_explanation_but_keeps_fg_chart(tmp_path: Path) -> None:
    pytest.importorskip("matplotlib")
    write_results(tmp_path, [risk_row("С1", 0, 1.0e-4, 1000.0)])

    charts = prepare_risk_charts(tmp_path)

    assert charts.fn_path is None
    assert charts.fg_path is not None
    document = Document()
    document.add_paragraph(FN_MARKER)
    assert render_risk_chart(document, FN_MARKER, None, FN_EMPTY_TEXT)
    assert document.paragraphs[0].text == FN_EMPTY_TEXT


def test_stale_fn_points_are_rejected(tmp_path: Path) -> None:
    rows = [risk_row("С1", 1, 1.0e-4, 1000.0)]
    write_results(tmp_path, rows)
    summary_path = tmp_path / "risk_summary.json"
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    summary["fn_points"][1]["cumulative_frequency"] = 2.0e-4
    write_json(summary_path, summary)

    with pytest.raises(ReportRiskChartsError, match="точки F/N не совпадают"):
        prepare_risk_charts(tmp_path)
