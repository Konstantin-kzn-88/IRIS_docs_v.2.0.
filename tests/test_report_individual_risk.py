import json
from pathlib import Path

import pytest
from docx import Document

from iris_v2.report_individual_risk import (
    ReportIndividualRiskError,
    load_individual_risk_section,
    render_individual_risk_section,
)


def write_json(path: Path, value: object) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")


def risk_row(
    code: str,
    component: str,
    fatalities: float | None,
    injured: float | None,
    people: int,
) -> dict:
    return {
        "scenario_code": code,
        "hazard_component": component,
        "risk_people_count": people,
        "individual_risk_status": "calculated" if people else "no_people",
        "individual_risk_fatalities": fatalities,
        "individual_risk_injured": injured,
    }


def test_individual_risk_is_grouped_and_totalled(tmp_path: Path) -> None:
    results = [
        risk_row("С1", "Трубопроводы", 1.0e-6, 2.0e-6, 20),
        risk_row("С2", "Резервуарный парк", 3.0e-6, 4.0e-6, 20),
        risk_row("С3", "Трубопроводы", 5.0e-6, 6.0e-6, 20),
    ]
    write_json(tmp_path / "risk_results.json", {"results": results})
    write_json(
        tmp_path / "risk_summary.json",
        {
            "case_count": 3,
            "component_count": 2,
            "risk_unit": "1/год",
            "total_individual_risk_fatalities": 9.0e-6,
            "total_individual_risk_injured": 1.2e-5,
            "components": [
                {
                    "hazard_component": "Трубопроводы",
                    "scenario_count": 2,
                    "individual_risk_fatalities": 6.0e-6,
                    "individual_risk_injured": 8.0e-6,
                },
                {
                    "hazard_component": "Резервуарный парк",
                    "scenario_count": 1,
                    "individual_risk_fatalities": 3.0e-6,
                    "individual_risk_injured": 4.0e-6,
                },
            ],
        },
    )

    result = load_individual_risk_section(tmp_path)

    assert not result.no_people
    assert result.rows[-1] == {
        "component": "Итого по ОПО",
        "fatalities": "9.000E-06",
        "injured": "1.200E-05",
    }


def test_zero_people_returns_explanation_mode(tmp_path: Path) -> None:
    write_json(
        tmp_path / "risk_results.json",
        {"results": [risk_row("С1", "Трубопроводы", None, None, 0)]},
    )
    write_json(
        tmp_path / "risk_summary.json",
        {
            "case_count": 1,
            "component_count": 1,
            "risk_unit": "1/год",
            "total_individual_risk_fatalities": None,
            "total_individual_risk_injured": None,
            "components": [
                {
                    "hazard_component": "Трубопроводы",
                    "scenario_count": 1,
                    "individual_risk_fatalities": None,
                    "individual_risk_injured": None,
                }
            ],
        },
    )

    result = load_individual_risk_section(tmp_path)

    assert result.no_people
    assert result.rows == ()

    document = Document()
    document.add_paragraph("{{INDIVIDUAL_RISK_SECTION}}")
    assert render_individual_risk_section(document, result)
    assert document.paragraphs[0].text == (
        "Индивидуальный риск не рассчитан: численность людей на ОПО и соседних "
        "ОПО равна нулю."
    )


def test_stale_individual_risk_is_rejected(tmp_path: Path) -> None:
    write_json(
        tmp_path / "risk_results.json",
        {"results": [risk_row("С1", "Трубопроводы", 1.0e-6, 2.0e-6, 20)]},
    )
    write_json(
        tmp_path / "risk_summary.json",
        {
            "case_count": 1,
            "component_count": 1,
            "risk_unit": "1/год",
            "total_individual_risk_fatalities": 2.0e-6,
            "total_individual_risk_injured": 2.0e-6,
            "components": [
                {
                    "hazard_component": "Трубопроводы",
                    "scenario_count": 1,
                    "individual_risk_fatalities": 1.0e-6,
                    "individual_risk_injured": 2.0e-6,
                }
            ],
        },
    )

    with pytest.raises(ReportIndividualRiskError, match="итоговый индивидуальный риск"):
        load_individual_risk_section(tmp_path)
