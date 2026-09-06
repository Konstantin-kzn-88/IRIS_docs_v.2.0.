import hashlib
import json
from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from iris_v2.project_common import ProjectCommonService, new_project_common
from iris_v2.report_generation import ReportGenerationError, ReportGenerationService
from iris_v2.service import CreateProjectData, ProjectService


def make_project(tmp_path: Path) -> Path:
    project = tmp_path / "project"
    ProjectService().create(
        project,
        CreateProjectData(
            name="Проект ДПБ",
            code="DPB-001",
            organization_name="АО Короткое",
            opo_name="Площадка нефти",
            opo_registration_number="А00-00001-0001",
            organization_snapshot={
                "organization": {
                    "full_name": "Акционерное общество Короткое",
                    "short_name": "АО Короткое",
                    "ids": {"inn": "1234567890"},
                    "contacts": {"phone": "+7 000 000-00-00"},
                },
                "permits": {"license_number": "Лицензия № 1"},
            },
            opo_snapshot={
                "site_id": "opo_0001",
                "name": "Площадка нефти",
                "reg_number": "А00-00001-0001",
                "object_id": "III",
                "sanitary_protection_zone_m": 0,
                "personnel": {
                    "employees_count": 15,
                    "employees_other_opo_count": 4,
                },
            },
        ),
    )
    common = new_project_common("Проект ДПБ", "DPB-001")
    common["dpb_code"] = "ДПБ-01"
    common["executor"]["name"] = "ООО Разработчик"
    ProjectCommonService().save(project, common)
    return project


def install_template(project: Path, unknown_marker: bool = False) -> Path:
    path = project / "input" / "templates" / "selected" / "template_report.docx"
    document = Document()
    paragraph = document.add_paragraph("Проект: ")
    paragraph.add_run("{{ PRO")
    paragraph.add_run("JECT_NAME }}")
    document.add_paragraph("Шифр: {{ DPB_CODE }}")
    document.add_paragraph("Дата: {{ generated_at }}")
    document.add_paragraph("СЗЗ: {{ SITE_SANITARY_PROTECTION_ZONE_M }}")
    document.add_paragraph("{{SUBSTANCES_SECTION}}")
    document.add_paragraph("{{EQUIPMENT_SECTION}}")
    document.add_paragraph("{{DISTRIBUTION_SECTION}}")
    document.add_paragraph("{{SCENARIOS_SECTION}}")
    document.add_paragraph("{{OV_AMOUNT_SECTION}}")
    document.add_paragraph("{{IMPACT_ZONES_SECTION}}")
    document.add_paragraph("{{CASUALTIES_SECTION}}")
    document.add_paragraph("{{DAMAGE_SECTION}}")
    document.add_paragraph("{{FATAL_ACCIDENT_FREQUENCY}}")
    document.add_paragraph("{{COLLECTIVE_RISK_SECTION}}")
    document.add_paragraph("{{INDIVIDUAL_RISK_SECTION}}")
    document.add_paragraph("{{MAX_DAMAGE_BY_COMPONENT_SECTION}}")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Организация: {{ FULL_NAME }}"
    document.sections[0].header.paragraphs[0].text = (
        "{{ SHORT_NAME }} — {{ SITE_NAME }}"
    )
    if unknown_marker:
        document.add_paragraph("{{ UNKNOWN_FIELD }}")
    document.save(path)
    config = {
        "format_version": 1,
        "template_profile": "test",
        "documents": [
            {
                "name": path.name,
                "path": "input/templates/selected/template_report.docx",
                "size": path.stat().st_size,
                "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
            }
        ],
    }
    (project / "report_config.json").write_text(
        json.dumps(config, ensure_ascii=False), encoding="utf-8"
    )
    return path


def write_substances(project: Path) -> None:
    data = [
        {
            "id": 1,
            "name": "Нефть",
            "kind": 0,
            "formula": "Смесь углеводородов",
            "composition": {
                "components": [
                    {"name": "Углеводороды", "mass_fraction": 0.98},
                    {"name": "Вода", "mass_fraction": 0.02},
                ]
            },
            "notes": "Товарная нефть",
            "physical": {
                "density_liquid_kg_per_m3": 850,
                "boiling_point_C": None,
            },
            "explosion": {"flash_point_C": -35},
            "toxicity": {"hazard_class": 3},
            "reactivity": "Стабильна при нормальных условиях",
            "odor": "Характерный",
            "corrosiveness": "",
            "precautions": "Исключить источники зажигания",
            "impact": "Опасна при попадании в окружающую среду",
            "protection": "Спецодежда и средства защиты органов дыхания",
            "neutralization_methods": "Сбор механическим способом",
            "first_aid": "Вывести пострадавшего на свежий воздух",
        },
    ]
    (project / "substances.json").write_text(
        json.dumps(data, ensure_ascii=False), encoding="utf-8"
    )


def write_equipment(project: Path) -> None:
    data = [
        {
            "id": 1,
            "source_id": 17,
            "substance_id": 1,
            "equipment_name": "Нефтепровод от скважины № 1",
            "equipment_type": 9,
            "phase_state": "ж.ф.",
            "total_length_m": 1000.0,
            "equipment_count": None,
            "accident_section_length_m": 100.0,
            "diameter_mm": 219.0,
            "wall_thickness_mm": 8.0,
            "volume_m3": None,
            "fill_fraction": 0.8,
            "pressure_mpa": 1.6,
            "spill_coefficient": 20.0,
            "spill_area_m2": 0.0,
            "substance_temperature_c": 20.0,
            "shutdown_time_s": 12.0,
            "evaporation_time_s": 3600.0,
            "hazard_component": "Участок трубопроводов",
            "clutter_degree": 2,
            "coord_type": 0,
            "coordinates": [10.5, 20.5, 30.0],
            "possible_dead": 2,
            "possible_injured": 4,
        },
        {
            "id": 2,
            "source_id": 18,
            "substance_id": 1,
            "equipment_name": "Резервуар нефти РВС-5000",
            "equipment_type": 1,
            "phase_state": "ж.ф.",
            "total_length_m": None,
            "equipment_count": 2,
            "accident_section_length_m": None,
            "diameter_mm": None,
            "wall_thickness_mm": None,
            "volume_m3": 5000.0,
            "fill_fraction": 0.85,
            "pressure_mpa": 0.1,
            "spill_coefficient": 20.0,
            "spill_area_m2": 1200.0,
            "substance_temperature_c": 15.0,
            "shutdown_time_s": 12.0,
            "evaporation_time_s": 3600.0,
            "hazard_component": "Резервуарный парк",
            "clutter_degree": 2,
            "coord_type": 1,
            "coordinates": [50.0, 60.0],
            "possible_dead": 3,
            "possible_injured": 7,
        },
    ]
    (project / "equipments.json").write_text(
        json.dumps(data, ensure_ascii=False), encoding="utf-8"
    )


def write_amount_results(project: Path) -> None:
    data = {
        "format_version": 1,
        "equipment_count": 2,
        "results": [
            {"equipment_id": 1, "amount_t": 5.0},
            {"equipment_id": 2, "amount_t": 10.123},
        ],
    }
    (project / "amount_results.json").write_text(
        json.dumps(data, ensure_ascii=False), encoding="utf-8"
    )


def write_scenario_results(project: Path) -> None:
    case = {
        "id": 1,
        "scenario_code": "С1",
        "equipment_id": 1,
        "equipment_name": "Нефтепровод от скважины № 1",
        "hazard_component": "Участок трубопроводов",
        "scenario_text": "Разрыв трубопровода → пожар пролива",
    }
    (project / "calculation_cases.json").write_text(
        json.dumps(
            {"format_version": 1, "case_count": 1, "cases": [case]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    frequency = dict(case)
    frequency["scenario_frequency"] = 6.0e-5
    (project / "frequency_results.json").write_text(
        json.dumps(
            {"format_version": 1, "case_count": 1, "results": [frequency]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    release = dict(case)
    release["ov_in_accident_t"] = 4.321
    (project / "release_results.json").write_text(
        json.dumps(
            {"format_version": 1, "case_count": 1, "results": [release]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    factor = dict(release)
    factor["calc_code"] = 1
    factor["ov_in_hazard_factor_t"] = 1.234
    (project / "hazard_factor_results.json").write_text(
        json.dumps(
            {"format_version": 1, "case_count": 1, "results": [factor]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    impact = dict(factor)
    impact["calc_code"] = 1
    impact["impact_values"] = {
        "q_10_5_m": 10.04,
        "q_7_0_m": 20.05,
        "q_4_2_m": 30.06,
        "q_1_4_m": 40.07,
    }
    impact["spill_area_m2"] = 125.55
    (project / "impact_zones.json").write_text(
        json.dumps(
            {"format_version": 1, "case_count": 1, "results": [impact]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    people = dict(impact)
    people.update(
        {
            "possible_dead": 2,
            "possible_injured": 4,
            "fatalities_count": 1,
            "injured_count": 3,
            "people_rule": "full_pool_fire",
        }
    )
    (project / "people_results.json").write_text(
        json.dumps(
            {"format_version": 1, "case_count": 1, "results": [people]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    damage = dict(people)
    damage.update(
        {
            "damage_unit": "тыс. руб.",
            "damage_scale": 30.3,
            "damage_scenario_coefficient": 0.8,
            "direct_losses": 100.0,
            "liquidation_costs": 10.0,
            "social_losses": 3250.0,
            "indirect_damage": 510.2,
            "total_environmental_damage": 23.6,
            "total_damage": 3893.8,
        }
    )
    (project / "damage_results.json").write_text(
        json.dumps(
            {"format_version": 1, "case_count": 1, "results": [damage]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    risk = dict(damage)
    risk.update(
        {
            "scenario_frequency": 6.0e-5,
            "risk_people_count": 15,
            "individual_risk_status": "calculated",
            "collective_risk_fatalities": 6.0e-5,
            "collective_risk_injured": 1.8e-4,
            "individual_risk_fatalities": 4.0e-6,
            "individual_risk_injured": 1.2e-5,
            "expected_damage": 0.233628,
        }
    )
    (project / "risk_results.json").write_text(
        json.dumps(
            {"format_version": 1, "case_count": 1, "results": [risk]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    (project / "risk_summary.json").write_text(
        json.dumps(
            {
                "format_version": 1,
                "case_count": 1,
                "component_count": 1,
                "fatal_accident_frequency_min": 6.0e-5,
                "fatal_accident_frequency_max": 6.0e-5,
                "total_collective_risk_fatalities": 6.0e-5,
                "total_collective_risk_injured": 1.8e-4,
                "total_individual_risk_fatalities": 4.0e-6,
                "total_individual_risk_injured": 1.2e-5,
                "risk_unit": "1/год",
                "damage_unit": "тыс. руб.",
                "components": [
                    {
                        "hazard_component": "Участок трубопроводов",
                        "scenario_count": 1,
                        "collective_risk_fatalities": 6.0e-5,
                        "collective_risk_injured": 1.8e-4,
                        "individual_risk_fatalities": 4.0e-6,
                        "individual_risk_injured": 1.2e-5,
                        "max_direct_losses": 100.0,
                        "max_total_environmental_damage": 23.6,
                        "max_total_damage": 3893.8,
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )


def all_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(paragraph.text for paragraph in cell.paragraphs)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
    return "\n".join(values)


def test_scalar_markers_are_filled_and_blocks_are_preserved(tmp_path: Path) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    write_equipment(project)
    write_amount_results(project)
    write_scenario_results(project)

    result = ReportGenerationService().generate(
        project, generated_at=datetime(2026, 9, 6, 12, 30)
    )

    text = all_text(Document(result.output_path))
    assert "Проект: Проект ДПБ" in text
    assert "Шифр: ДПБ-01" in text
    assert "Дата: 06.09.2026 12:30" in text
    assert "СЗЗ: отсутствует" in text
    assert "Организация: Акционерное общество Короткое" in text
    assert "АО Короткое — Площадка нефти" in text
    assert "{{SUBSTANCES_SECTION}}" not in text
    assert "Параметр" in text
    assert "Легковоспламеняющаяся жидкость (ЛВЖ)" not in text
    assert "Вид опасного вещества" not in text
    assert "Плотность жидкости" in text
    assert "850 кг/м³" in text
    assert "Температура кипения" not in text
    assert "{{EQUIPMENT_SECTION}}" not in text
    assert "Нефтепровод от скважины № 1" in text
    assert "Общая протяжённость для расчёта частоты" in text
    assert "1000 м" in text
    assert "Длина аварийного участка" in text
    assert "100 м" in text
    assert "Наружный диаметр" in text
    assert "219 мм" in text
    assert "Опасное вещество" in text
    assert "Нефть" in text
    assert "Резервуар нефти РВС-5000" in text
    assert "Количество оборудования для расчёта частоты" in text
    assert "2 шт." in text
    assert "Объём оборудования" in text
    assert "5000 м³" in text
    assert "Степень заполнения" in text
    assert "0,85" in text
    assert "Координаты" not in text
    assert "Возможные погибшие" not in text
    assert "{{DISTRIBUTION_SECTION}}" not in text
    assert "Технологический блок, оборудование" in text
    assert "Нефтепровод от скважины № 1 (Нефть)" not in text
    assert "10,123" in text
    assert "20,246" in text
    assert "Общая масса опасных веществ в оборудовании: 25,246 т" in text
    assert "{{SCENARIOS_SECTION}}" not in text
    assert "Разрыв трубопровода → пожар пролива" in text
    assert "Нефтепровод от скважины № 1 (Участок трубопроводов)" in text
    assert "6.000E-05" in text
    assert "{{OV_AMOUNT_SECTION}}" not in text
    assert "Количество ОВ, участвующего в аварии, т" in text
    assert "4,321" in text
    assert "1,234" in text
    assert "{{IMPACT_ZONES_SECTION}}" not in text
    assert "q=10,5" in text
    assert "10,0" in text
    assert "125,5" in text
    assert "интенсивность теплового излучения" in text
    assert "{{CASUALTIES_SECTION}}" not in text
    assert "Количество погибших, чел." in text
    assert "Количество пострадавших, чел." in text
    casualties_table = next(
        table
        for table in Document(result.output_path).tables
        if table.cell(0, 2).text == "Количество погибших, чел."
    )
    assert [cell.text for cell in casualties_table.rows[1].cells] == [
        "С1",
        "Нефтепровод от скважины № 1 (Участок трубопроводов)",
        "1",
        "3",
    ]
    assert "{{DAMAGE_SECTION}}" not in text
    assert "Затраты на ЛЛА" in text
    assert "ЛЛА — локализация и ликвидация аварии" in text
    assert "Все стоимостные показатели приведены в тыс. руб." in text
    damage_table = next(
        table
        for table in Document(result.output_path).tables
        if table.cell(0, 2).text == "Прямые потери"
    )
    assert [cell.text for cell in damage_table.rows[1].cells] == [
        "С1",
        "Нефтепровод от скважины № 1 (Участок трубопроводов)",
        "100,0",
        "10,0",
        "3250,0",
        "510,2",
        "23,6",
        "3893,8",
    ]
    assert "{{FATAL_ACCIDENT_FREQUENCY}}" not in text
    assert "Частота сценариев с погибшими: 6.000E-05 1/год." in text
    assert "{{COLLECTIVE_RISK_SECTION}}" not in text
    collective_table = next(
        table
        for table in Document(result.output_path).tables
        if table.cell(0, 0).text == "Составляющая ОПО"
    )
    assert [cell.text for cell in collective_table.rows[1].cells] == [
        "Участок трубопроводов",
        "6.000E-05",
        "1.800E-04",
    ]
    assert [cell.text for cell in collective_table.rows[2].cells] == [
        "Итого по ОПО",
        "6.000E-05",
        "1.800E-04",
    ]
    assert "{{INDIVIDUAL_RISK_SECTION}}" not in text
    individual_table = next(
        table
        for table in Document(result.output_path).tables
        if table.cell(0, 1).text == "Индивидуальный риск гибели, 1/год"
    )
    assert [cell.text for cell in individual_table.rows[1].cells] == [
        "Участок трубопроводов",
        "4.000E-06",
        "1.200E-05",
    ]
    assert [cell.text for cell in individual_table.rows[2].cells] == [
        "Итого по ОПО",
        "4.000E-06",
        "1.200E-05",
    ]
    assert "{{MAX_DAMAGE_BY_COMPONENT_SECTION}}" not in text
    maximum_damage_table = next(
        table
        for table in Document(result.output_path).tables
        if table.cell(0, 1).text == "Максимальные прямые потери, тыс. руб."
    )
    assert [cell.text for cell in maximum_damage_table.rows[1].cells] == [
        "Участок трубопроводов",
        "100,0",
        "23,6",
        "3893,8",
    ]
    assert [cell.text for cell in maximum_damage_table.rows[2].cells] == [
        "Максимум по ОПО",
        "100,0",
        "23,6",
        "3893,8",
    ]
    assert all(
        row._tr.get_or_add_trPr().find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit"
        ) is not None
        for table in Document(result.output_path).tables[:2]
        for row in table.rows
    )
    assert result.replaced_count == 7
    assert result.filled_sections == (
        "SUBSTANCES_SECTION",
        "EQUIPMENT_SECTION",
        "DISTRIBUTION_SECTION",
        "SCENARIOS_SECTION",
        "OV_AMOUNT_SECTION",
        "IMPACT_ZONES_SECTION",
        "CASUALTIES_SECTION",
        "DAMAGE_SECTION",
        "FATAL_ACCIDENT_FREQUENCY",
        "COLLECTIVE_RISK_SECTION",
        "INDIVIDUAL_RISK_SECTION",
        "MAX_DAMAGE_BY_COMPONENT_SECTION",
    )
    assert result.deferred_markers == ()


def test_unknown_marker_does_not_replace_existing_report(tmp_path: Path) -> None:
    project = make_project(tmp_path)
    install_template(project, unknown_marker=True)
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="UNKNOWN_FIELD"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"


def test_changed_template_is_rejected(tmp_path: Path) -> None:
    project = make_project(tmp_path)
    template = install_template(project)
    template.write_bytes(template.read_bytes() + b"changed")

    with pytest.raises(ReportGenerationError, match="изменён после выбора"):
        ReportGenerationService().generate(project)


def test_missing_equipment_does_not_replace_existing_report(tmp_path: Path) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="не импортировано"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"


def test_builtin_default_template_contains_only_supported_markers(
    tmp_path: Path,
    monkeypatch,
) -> None:
    monkeypatch.chdir(tmp_path)
    project = make_project(tmp_path)
    write_substances(project)
    write_equipment(project)
    write_amount_results(project)
    write_scenario_results(project)

    result = ReportGenerationService().generate(project)

    assert result.output_path.is_file()
    assert result.replaced_count > 0
    assert result.filled_sections == (
        "SUBSTANCES_SECTION",
        "EQUIPMENT_SECTION",
        "DISTRIBUTION_SECTION",
        "SCENARIOS_SECTION",
        "OV_AMOUNT_SECTION",
        "IMPACT_ZONES_SECTION",
        "CASUALTIES_SECTION",
        "DAMAGE_SECTION",
        "FATAL_ACCIDENT_FREQUENCY",
        "COLLECTIVE_RISK_SECTION",
        "INDIVIDUAL_RISK_SECTION",
        "MAX_DAMAGE_BY_COMPONENT_SECTION",
    )
    assert "SUBSTANCES_SECTION" not in result.deferred_markers
    assert "EQUIPMENT_SECTION" not in result.deferred_markers
    assert "DISTRIBUTION_SECTION" not in result.deferred_markers
    assert "SCENARIOS_SECTION" not in result.deferred_markers
    assert "OV_AMOUNT_SECTION" not in result.deferred_markers
    assert "IMPACT_ZONES_SECTION" not in result.deferred_markers
    assert "CASUALTIES_SECTION" not in result.deferred_markers
    assert "DAMAGE_SECTION" not in result.deferred_markers
    assert "FATAL_ACCIDENT_FREQUENCY" not in result.deferred_markers
    assert "COLLECTIVE_RISK_SECTION" not in result.deferred_markers
    assert "INDIVIDUAL_RISK_SECTION" not in result.deferred_markers
    assert "MAX_DAMAGE_BY_COMPONENT_SECTION" not in result.deferred_markers


def test_missing_amount_results_does_not_replace_existing_report(
    tmp_path: Path,
) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    write_equipment(project)
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="Количество опасного вещества"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"


def test_missing_frequency_results_does_not_replace_existing_report(
    tmp_path: Path,
) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    write_equipment(project)
    write_amount_results(project)
    write_scenario_results(project)
    (project / "frequency_results.json").unlink()
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="Частоты сценариев"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"


def test_missing_hazard_factor_results_does_not_replace_existing_report(
    tmp_path: Path,
) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    write_equipment(project)
    write_amount_results(project)
    write_scenario_results(project)
    (project / "hazard_factor_results.json").unlink()
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="поражающем факторе"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"


def test_missing_impact_zones_does_not_replace_existing_report(
    tmp_path: Path,
) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    write_equipment(project)
    write_amount_results(project)
    write_scenario_results(project)
    (project / "impact_zones.json").unlink()
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="Зоны поражающих факторов"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"


def test_missing_people_results_does_not_replace_existing_report(
    tmp_path: Path,
) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    write_equipment(project)
    write_amount_results(project)
    write_scenario_results(project)
    (project / "people_results.json").unlink()
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="Число погибших"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"


def test_missing_damage_results_does_not_replace_existing_report(
    tmp_path: Path,
) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    write_equipment(project)
    write_amount_results(project)
    write_scenario_results(project)
    (project / "damage_results.json").unlink()
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="Ущерб не рассчитан"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"


def test_missing_risk_summary_does_not_replace_existing_report(
    tmp_path: Path,
) -> None:
    project = make_project(tmp_path)
    install_template(project)
    write_substances(project)
    write_equipment(project)
    write_amount_results(project)
    write_scenario_results(project)
    (project / "risk_summary.json").unlink()
    output = project / "output" / "template_report_out.docx"
    output.write_bytes(b"previous report")

    with pytest.raises(ReportGenerationError, match="Сводные показатели риска"):
        ReportGenerationService().generate(project)

    assert output.read_bytes() == b"previous report"
