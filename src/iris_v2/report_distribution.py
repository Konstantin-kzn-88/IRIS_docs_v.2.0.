import json
import math
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.amount_calculation import FILE_NAME
from iris_v2.equipment import PIPELINE_TYPES


MARKER = "{{DISTRIBUTION_SECTION}}"


class ReportDistributionError(Exception):
    pass


def load_amount_results(
    project_directory: Path | str,
) -> tuple[dict[str, Any], ...]:
    path = Path(project_directory) / FILE_NAME
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportDistributionError(
            "Количество опасного вещества не рассчитано. "
            "Сначала выполните модуль «Количество ОВ»"
        ) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportDistributionError(
            f"Не удалось прочитать {FILE_NAME}"
        ) from exc
    values = raw.get("results") if isinstance(raw, dict) else None
    if not isinstance(values, list) or not values:
        raise ReportDistributionError(
            f"Файл {FILE_NAME} повреждён: results должен быть непустым списком"
        )
    if any(not isinstance(item, dict) for item in values):
        raise ReportDistributionError(
            f"Файл {FILE_NAME} повреждён: results содержит запись неверного формата"
        )
    return tuple(values)


def _number(value: float, digits: int) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def _finite_number(value: Any, label: str, *, minimum: float = 0.0) -> float:
    if (
        isinstance(value, bool)
        or not isinstance(value, (int, float))
        or not math.isfinite(float(value))
        or float(value) < minimum
    ):
        raise ReportDistributionError(
            f"{label} должно быть конечным числом не меньше {minimum:g}"
        )
    return float(value)


def distribution_rows(
    equipment: tuple[dict[str, Any], ...],
    substances: tuple[dict[str, Any], ...],
    amounts: tuple[dict[str, Any], ...],
) -> tuple[tuple[dict[str, str], ...], float]:
    substance_names = {
        item.get("id"): str(item.get("name", "")).strip()
        for item in substances
        if isinstance(item.get("id"), int) and str(item.get("name", "")).strip()
    }
    amount_by_equipment: dict[int, dict[str, Any]] = {}
    for index, item in enumerate(amounts, start=1):
        equipment_id = item.get("equipment_id")
        if (
            isinstance(equipment_id, bool)
            or not isinstance(equipment_id, int)
            or equipment_id <= 0
            or equipment_id in amount_by_equipment
        ):
            raise ReportDistributionError(
                f"{FILE_NAME}, запись {index}: недопустимый или повторяющийся equipment_id"
            )
        amount_by_equipment[equipment_id] = item

    grouped: dict[str, list[dict[str, str]]] = {}
    total_mass_t = 0.0
    equipment_ids: set[int] = set()
    for index, item in enumerate(equipment, start=1):
        equipment_id = item.get("id")
        if isinstance(equipment_id, bool) or not isinstance(equipment_id, int):
            raise ReportDistributionError(
                f"Оборудование {index}: недопустимый id"
            )
        if equipment_id in equipment_ids:
            raise ReportDistributionError(
                f"Оборудование {index}: повторяющийся id={equipment_id}"
            )
        equipment_ids.add(equipment_id)
        amount = amount_by_equipment.get(equipment_id)
        if amount is None:
            raise ReportDistributionError(
                f"В {FILE_NAME} отсутствует оборудование id={equipment_id}"
            )
        substance_id = item.get("substance_id")
        substance_name = substance_names.get(substance_id)
        if substance_name is None:
            raise ReportDistributionError(
                f"Оборудование {equipment_id}: substance_id отсутствует в substances.json"
            )
        equipment_name = str(item.get("equipment_name", "")).strip()
        component = str(item.get("hazard_component", "")).strip()
        if not equipment_name or not component:
            raise ReportDistributionError(
                f"Оборудование {equipment_id}: не заполнено наименование или составляющая ОПО"
            )

        mass_in_unit = _finite_number(
            amount.get("amount_t"),
            f"{FILE_NAME}, оборудование {equipment_id}: amount_t",
        )
        equipment_type = item.get("equipment_type")
        if equipment_type in PIPELINE_TYPES:
            units = 1
        else:
            units_value = _finite_number(
                item.get("equipment_count"),
                f"Оборудование {equipment_id}: equipment_count",
                minimum=1.0,
            )
            if not units_value.is_integer():
                raise ReportDistributionError(
                    f"Оборудование {equipment_id}: equipment_count должно быть целым числом"
                )
            units = int(units_value)
        mass_in_block = mass_in_unit * units
        total_mass_t += mass_in_block

        pressure = _finite_number(
            item.get("pressure_mpa"),
            f"Оборудование {equipment_id}: pressure_mpa",
        )
        temperature = _finite_number(
            item.get("substance_temperature_c"),
            f"Оборудование {equipment_id}: substance_temperature_c",
            minimum=-273.15,
        )
        row = {
            "component": component,
            "equipment": f"{equipment_name} ({substance_name})",
            "units": str(units),
            "mass_in_unit": _number(mass_in_unit, 3),
            "mass_in_block": _number(mass_in_block, 3),
            "state": str(item.get("phase_state", "")).strip() or "—",
            "pressure": _number(pressure, 2),
            "temperature": _number(temperature, 1),
        }
        grouped.setdefault(component, []).append(row)

    extra_ids = set(amount_by_equipment) - equipment_ids
    if extra_ids:
        values = ", ".join(str(value) for value in sorted(extra_ids))
        raise ReportDistributionError(
            f"В {FILE_NAME} найдено отсутствующее в equipments.json оборудование: {values}"
        )

    rows: list[dict[str, str]] = []
    for component_rows in grouped.values():
        rows.extend(sorted(component_rows, key=lambda value: value["equipment"]))
    return tuple(rows), total_mass_t


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_table_geometry(document: DocumentType, table: Any) -> None:
    section = document.sections[0]
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    proportions = (0.18, 0.27, 0.07, 0.10, 0.10, 0.08, 0.09, 0.11)
    widths = [int(total_twips * value) for value in proportions[:-1]]
    widths.append(total_twips - sum(widths))
    table.autofit = False

    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        cells = row._tr.tc_lst
        if len(cells) == 8:
            row_widths = widths
        elif len(cells) == 3:
            row_widths = (
                sum(widths[:3]),
                sum(widths[3:5]),
                sum(widths[5:]),
            )
        else:
            continue
        for cell, width in zip(cells, row_widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_distribution_section(
    document: DocumentType,
    equipment: tuple[dict[str, Any], ...],
    substances: tuple[dict[str, Any], ...],
    amounts: tuple[dict[str, Any], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False
    rows, total_mass_t = distribution_rows(equipment, substances, amounts)

    table = document.add_table(rows=2, cols=8)
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)

    group_header = table.rows[0].cells
    equipment_header = group_header[0].merge(group_header[2])
    amount_header = group_header[3].merge(group_header[4])
    conditions_header = group_header[5].merge(group_header[7])
    for cell, value in (
        (equipment_header, "Технологический блок, оборудование"),
        (amount_header, "Количество опасного вещества, т"),
        (conditions_header, "Физические условия содержания опасного вещества"),
    ):
        _set_cell_text(cell, value, bold=True, centered=True)
        _shade(cell, "D9E1F2")

    headers = (
        "Наименование составляющей",
        "Наименование оборудования, № по схеме (опасное вещество)",
        "Кол-во единиц",
        "В единице оборудования",
        "В блоке",
        "Агр. состояние",
        "Давление, МПа",
        "Температура, °C",
    )
    for cell, value in zip(table.rows[1].cells, headers):
        _set_cell_text(cell, value, bold=True, centered=True)
        _shade(cell, "D9E1F2")

    for item in rows:
        cells = table.add_row().cells
        values = (
            item["component"],
            item["equipment"],
            item["units"],
            item["mass_in_unit"],
            item["mass_in_block"],
            item["state"],
            item["pressure"],
            item["temperature"],
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            _set_cell_text(cell, value, centered=index >= 2)

    total_paragraph = document.add_paragraph()
    total_paragraph.paragraph_format.space_before = Pt(6)
    total_paragraph.paragraph_format.space_after = Pt(0)
    total_run = total_paragraph.add_run(
        f"Общая масса опасных веществ в оборудовании: "
        f"{_number(total_mass_t, 3)} т"
    )
    total_run.font.name = "Times New Roman"
    total_run.font.size = Pt(11)
    total_fonts = total_run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        total_fonts.set(qn(f"w:{name}"), "Times New Roman")
    table._tbl.addnext(total_paragraph._p)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    _set_repeat_header(table.rows[0])
    _set_repeat_header(table.rows[1])
    for row in table.rows:
        _prevent_row_split(row)
    _set_table_geometry(document, table)
    return True
