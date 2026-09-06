import json
import math
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from iris_v2.risk_calculation import FILE_NAME as RISK_FILE_NAME
from iris_v2.risk_charts import (
    FG_FILE_NAME,
    FN_FILE_NAME,
    _save_fg_chart,
    _save_fn_chart,
)
from iris_v2.risk_summary import (
    FILE_NAME as SUMMARY_FILE_NAME,
    build_fg_points,
    build_fn_points,
)


FN_MARKER = "{{FN_CHART}}"
FG_MARKER = "{{FG_CHART}}"
FN_EMPTY_TEXT = "F/N-диаграмма не построена: сценарии с погибшими отсутствуют."
FG_EMPTY_TEXT = "F/G-диаграмма не построена: сценарии с положительным ущербом отсутствуют."


class ReportRiskChartsError(Exception):
    pass


@dataclass(frozen=True)
class ReportRiskCharts:
    fn_path: Path | None
    fg_path: Path | None


def _load_object(path: Path, missing_message: str) -> dict[str, Any]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportRiskChartsError(missing_message) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportRiskChartsError(f"Не удалось прочитать {path.name}") from exc
    if not isinstance(raw, dict):
        raise ReportRiskChartsError(f"Файл {path.name} повреждён: ожидается объект")
    return raw


def _number(value: Any, label: str) -> float:
    if (
        isinstance(value, bool)
        or not isinstance(value, (int, float))
        or not math.isfinite(float(value))
        or float(value) < 0
    ):
        raise ReportRiskChartsError(f"{label} должно быть числом не меньше нуля")
    return float(value)


def _fn_points(value: Any) -> list[tuple[int, float]]:
    if not isinstance(value, list):
        raise ReportRiskChartsError(
            f"Файл {SUMMARY_FILE_NAME} не содержит точек F/N"
        )
    result: list[tuple[int, float]] = []
    for index, point in enumerate(value, start=1):
        if not isinstance(point, dict):
            raise ReportRiskChartsError(f"Точка F/N {index}: ожидается объект")
        fatalities = point.get("fatalities_count")
        if isinstance(fatalities, bool) or not isinstance(fatalities, int) or fatalities < 0:
            raise ReportRiskChartsError(
                f"Точка F/N {index}: fatalities_count должно быть целым числом "
                "не меньше нуля"
            )
        frequency = _number(
            point.get("cumulative_frequency"),
            f"Точка F/N {index}: cumulative_frequency",
        )
        if frequency <= 0:
            raise ReportRiskChartsError(
                f"Точка F/N {index}: частота должна быть больше нуля"
            )
        result.append((fatalities, frequency))
    return result


def _fg_points(value: Any) -> list[tuple[float, float]]:
    if not isinstance(value, list):
        raise ReportRiskChartsError(
            f"Файл {SUMMARY_FILE_NAME} не содержит точек F/G"
        )
    result: list[tuple[float, float]] = []
    for index, point in enumerate(value, start=1):
        if not isinstance(point, dict):
            raise ReportRiskChartsError(f"Точка F/G {index}: ожидается объект")
        damage = _number(
            point.get("damage_million_rub"),
            f"Точка F/G {index}: damage_million_rub",
        )
        frequency = _number(
            point.get("cumulative_frequency"),
            f"Точка F/G {index}: cumulative_frequency",
        )
        if frequency <= 0:
            raise ReportRiskChartsError(
                f"Точка F/G {index}: частота должна быть больше нуля"
            )
        result.append((damage, frequency))
    return result


def _points_match(
    stored: list[tuple[int | float, float]],
    expected: list[tuple[int | float, float]],
) -> bool:
    return len(stored) == len(expected) and all(
        math.isclose(float(left_x), float(right_x), rel_tol=1e-12, abs_tol=1e-12)
        and math.isclose(left_y, right_y, rel_tol=1e-12, abs_tol=0.0)
        for (left_x, left_y), (right_x, right_y) in zip(stored, expected)
    )


def prepare_risk_charts(project_directory: Path | str) -> ReportRiskCharts:
    project = Path(project_directory)
    risk = _load_object(
        project / RISK_FILE_NAME,
        "Риски не рассчитаны. Сначала выполните модуль «Расчёт риска»",
    )
    summary = _load_object(
        project / SUMMARY_FILE_NAME,
        "Сводные показатели риска не рассчитаны. "
        "Сначала выполните модуль «Сводные показатели риска»",
    )
    results = risk.get("results")
    if (
        not isinstance(results, list)
        or not results
        or any(not isinstance(item, dict) for item in results)
    ):
        raise ReportRiskChartsError(
            f"Файл {RISK_FILE_NAME} повреждён: results должен быть непустым "
            "списком объектов"
        )
    if summary.get("case_count") != len(results):
        raise ReportRiskChartsError(
            f"Файл {SUMMARY_FILE_NAME} устарел: не совпадает case_count"
        )
    if summary.get("risk_unit") != "1/год":
        raise ReportRiskChartsError(
            f"Файл {SUMMARY_FILE_NAME} содержит неизвестную единицу частоты"
        )
    if summary.get("fg_damage_unit") != "млн руб.":
        raise ReportRiskChartsError(
            f"Файл {SUMMARY_FILE_NAME} содержит неизвестную единицу ущерба F/G"
        )
    try:
        expected_fn_data = build_fn_points(results)
        expected_fg_data = build_fg_points(results)
    except ValueError as exc:
        raise ReportRiskChartsError(f"Файл {RISK_FILE_NAME} повреждён: {exc}") from exc
    expected_fn = [
        (int(point["fatalities_count"]), float(point["cumulative_frequency"]))
        for point in expected_fn_data
    ]
    expected_fg = [
        (float(point["damage_million_rub"]), float(point["cumulative_frequency"]))
        for point in expected_fg_data
    ]
    stored_fn = _fn_points(summary.get("fn_points"))
    stored_fg = _fg_points(summary.get("fg_points"))
    if not _points_match(stored_fn, expected_fn):
        raise ReportRiskChartsError(
            f"Файл {SUMMARY_FILE_NAME} устарел: точки F/N не совпадают с "
            f"{RISK_FILE_NAME}"
        )
    if not _points_match(stored_fg, expected_fg):
        raise ReportRiskChartsError(
            f"Файл {SUMMARY_FILE_NAME} устарел: точки F/G не совпадают с "
            f"{RISK_FILE_NAME}"
        )

    has_fn = any(fatalities > 0 for fatalities, _ in stored_fn)
    has_fg = any(damage > 0 for damage, _ in stored_fg)
    output_directory = project / "output" / "charts"
    fn_path = output_directory / FN_FILE_NAME if has_fn else None
    fg_path = output_directory / FG_FILE_NAME if has_fg else None
    fn_temporary = output_directory / f".{FN_FILE_NAME}.report.tmp"
    fg_temporary = output_directory / f".{FG_FILE_NAME}.report.tmp"
    try:
        if has_fn or has_fg:
            import matplotlib

            matplotlib.use("Agg")
            output_directory.mkdir(parents=True, exist_ok=True)
        if has_fn:
            _save_fn_chart(stored_fn, fn_temporary)
        if has_fg:
            _save_fg_chart(stored_fg, fg_temporary)
        if has_fn:
            fn_temporary.replace(fn_path)
        if has_fg:
            fg_temporary.replace(fg_path)
    except ImportError as exc:
        raise ReportRiskChartsError(
            "Не установлен matplotlib. Выполните: python -m pip install -e ."
        ) from exc
    except OSError as exc:
        raise ReportRiskChartsError("Не удалось сохранить диаграммы риска") from exc
    except Exception as exc:
        raise ReportRiskChartsError(
            f"Не удалось построить диаграммы риска: {exc}"
        ) from exc
    finally:
        fn_temporary.unlink(missing_ok=True)
        fg_temporary.unlink(missing_ok=True)
    return ReportRiskCharts(fn_path=fn_path, fg_path=fg_path)


def _set_font(run: Any, size: float) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")


def _paragraph_section(document: DocumentType, element: Any) -> Any:
    section_index = 0
    for child in document.element.body:
        if child is element:
            break
        if child.tag == qn("w:p") and child.find(
            f"./{qn('w:pPr')}/{qn('w:sectPr')}"
        ) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def render_risk_chart(
    document: DocumentType,
    marker: str,
    path: Path | None,
    empty_text: str,
) -> bool:
    paragraph = next(
        (item for item in document.paragraphs if marker in item.text),
        None,
    )
    if paragraph is None:
        return False
    paragraph.clear()
    if path is None:
        _set_font(paragraph.add_run(empty_text), 12)
        return True
    section = _paragraph_section(document, paragraph._p)
    available_width = section.page_width - section.left_margin - section.right_margin
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=min(available_width, Inches(6.5)))
    return True
