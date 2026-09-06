from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

MARKER = "{{SUBSTANCES_SECTION}}"

NESTED_FIELDS = {
    "composition": (
        "Состав",
        {
            "notes": ("Примечания", ""),
            "components": ("Компоненты", ""),
        },
    ),
    "physical": (
        "Физические свойства",
        {
            "molar_mass_kg_per_mol": ("Молярная масса", "кг/моль"),
            "density_liquid_kg_per_m3": ("Плотность жидкости", "кг/м³"),
            "density_gas_kg_per_m3": ("Плотность газа", "кг/м³"),
            "evaporation_heat_J_per_kg": ("Теплота испарения", "Дж/кг"),
            "boiling_point_C": ("Температура кипения", "°C"),
        },
    ),
    "explosion": (
        "Взрыво- и пожароопасность",
        {
            "explosion_hazard_class": ("Класс взрывоопасности", ""),
            "flash_point_C": ("Температура вспышки", "°C"),
            "lel_percent": ("НКПР", "%"),
            "autoignition_temp_C": ("Температура самовоспламенения", "°C"),
            "energy_reserve_factor": ("Коэффициент энергетического запаса", ""),
            "expansion_degree": ("Степень расширения", ""),
            "heat_of_combustion_kJ_per_kg": ("Теплота сгорания", "кДж/кг"),
            "burning_rate_kg_per_s_m2": ("Скорость горения", "кг/(с·м²)"),
        },
    ),
    "toxicity": (
        "Токсическая опасность",
        {
            "hazard_class": ("Класс опасности", ""),
            "pdk_mg_per_m3": ("ПДК", "мг/м³"),
            "lethal_tox_dose_mg_min_per_L": (
                "Смертельная токсодоза",
                "мг·мин/л",
            ),
            "threshold_tox_dose_mg_min_per_L": (
                "Пороговая токсодоза",
                "мг·мин/л",
            ),
        },
    ),
}

ADDITIONAL_FIELDS = (
    ("reactivity", "Реакционная способность"),
    ("odor", "Запах"),
    ("corrosiveness", "Коррозионная активность"),
    ("precautions", "Меры предосторожности"),
    (
        "impact",
        "Воздействие на людей и окружающую среду, в том числе от поражающих факторов аварии",
    ),
    ("protection", "Средства индивидуальной и коллективной защиты"),
    ("neutralization_methods", "Методы перевода вещества в безвредное состояние"),
    (
        "first_aid",
        "Меры первой помощи пострадавшим от воздействия поражающих факторов при аварии",
    ),
)


class ReportSubstancesError(Exception):
    pass


def _is_empty(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip() in ("", "-")
    if isinstance(value, (dict, list, tuple)):
        return not value
    return False


def _number(value: int | float) -> str:
    return format(value, ".12g").replace(".", ",")


def _value(value: Any, unit: str = "") -> str:
    if isinstance(value, bool):
        text = "да" if value else "нет"
    elif isinstance(value, (int, float)):
        text = _number(value)
    else:
        text = str(value).strip()
    return f"{text} {unit}".strip()


def _components(value: Any) -> str:
    if not isinstance(value, list):
        return _value(value)
    parts: list[str] = []
    for component in value:
        if not isinstance(component, dict):
            if not _is_empty(component):
                parts.append(_value(component))
            continue
        name = _value(component.get("name", "Компонент"))
        fraction = component.get("mass_fraction")
        if _is_empty(fraction):
            parts.append(name)
        elif isinstance(fraction, (int, float)):
            parts.append(f"{name} — {_number(fraction * 100)} %")
        else:
            parts.append(f"{name} — {_value(fraction)}")
    return "; ".join(parts)


def _nested_rows(
    data: Any,
    labels: dict[str, tuple[str, str]],
) -> list[tuple[str, str]]:
    if not isinstance(data, dict):
        return [] if _is_empty(data) else [("Значение", _value(data))]
    rows: list[tuple[str, str]] = []
    keys = list(labels) + sorted(key for key in data if key not in labels)
    for key in keys:
        value = data.get(key)
        if _is_empty(value):
            continue
        label, unit = labels.get(key, (key, ""))
        text = _components(value) if key == "components" else _value(value, unit)
        if text:
            rows.append((label, text))
    return rows


def substance_sections(item: dict[str, Any]) -> list[tuple[str, list[tuple[str, str]]]]:
    basic = [
        ("Наименование вещества", _value(item.get("name", "—"))),
    ]
    if not _is_empty(item.get("formula")):
        basic.append(("Химическая формула", _value(item["formula"])))
    if not _is_empty(item.get("notes")):
        basic.append(("Примечания", _value(item["notes"])))

    result = [("Основные сведения", basic)]
    for key, (title, labels) in NESTED_FIELDS.items():
        rows = _nested_rows(item.get(key), labels)
        if rows:
            result.append((title, rows))
    additional = [
        (label, _value(item[key]))
        for key, label in ADDITIONAL_FIELDS
        if not _is_empty(item.get(key))
    ]
    if additional:
        result.append(("Дополнительная информация", additional))
    return result


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _cell_margins(cell: Any, value: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), "Times New Roman"
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _set_table_geometry(document: DocumentType, table: Any) -> None:
    section = document.sections[0]
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    indent_twips = 100
    total_twips -= indent_twips
    widths = (int(total_twips * 0.36), total_twips - int(total_twips * 0.36))
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(indent_twips))
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        cells = row._tr.tc_lst
        row_widths = (total_twips,) if len(cells) == 1 else widths
        for cell, width in zip(cells, row_widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_substances_section(
    document: DocumentType,
    substances: tuple[dict[str, Any], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False
    if not substances:
        raise ReportSubstancesError(
            "Вещества проекта не выбраны. Сначала заполните модуль «Вещества»"
        )

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    header = table.rows[0].cells
    _set_cell_text(header[0], "Параметр", bold=True, centered=True)
    _set_cell_text(header[1], "Значение", bold=True, centered=True)
    _shade(header[0], "D9E1F2")
    _shade(header[1], "D9E1F2")
    _set_repeat_header(table.rows[0])

    for item in substances:
        title_cells = table.add_row().cells
        title_cell = title_cells[0].merge(title_cells[1])
        _set_cell_text(title_cell, _value(item.get("name", "—")), bold=True)
        _shade(title_cell, "EDEDED")
        for section_title, rows in substance_sections(item):
            section_cells = table.add_row().cells
            section_cell = section_cells[0].merge(section_cells[1])
            _set_cell_text(section_cell, section_title, bold=True)
            _shade(section_cell, "F5F5F5")
            for label, value in rows:
                cells = table.add_row().cells
                _set_cell_text(cells[0], label)
                _set_cell_text(cells[1], value)

    for row in table.rows:
        _prevent_row_split(row)
    _set_table_geometry(document, table)
    return True
