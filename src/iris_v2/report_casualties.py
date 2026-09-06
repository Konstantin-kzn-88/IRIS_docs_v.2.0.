import json
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.impact_zones import FILE_NAME as IMPACT_ZONES_FILE_NAME
from iris_v2.people_calculation import FILE_NAME as PEOPLE_FILE_NAME


MARKER = "{{CASUALTIES_SECTION}}"


class ReportCasualtiesError(Exception):
    pass


def _load_results(path: Path, missing_message: str) -> tuple[dict[str, Any], ...]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportCasualtiesError(missing_message) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportCasualtiesError(f"Не удалось прочитать {path.name}") from exc
    values = raw.get("results") if isinstance(raw, dict) else None
    if not isinstance(values, list) or not values:
        raise ReportCasualtiesError(
            f"Файл {path.name} повреждён: results должен быть непустым списком"
        )
    if any(not isinstance(item, dict) for item in values):
        raise ReportCasualtiesError(
            f"Файл {path.name} повреждён: results содержит запись неверного формата"
        )
    return tuple(values)


def _load_equipment(path: Path) -> dict[int, dict[str, Any]]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportCasualtiesError(
            "Оборудование не импортировано. Сначала импортируйте equipments.json"
        ) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportCasualtiesError("Не удалось прочитать equipments.json") from exc
    if not isinstance(raw, list) or not raw:
        raise ReportCasualtiesError(
            "Файл equipments.json повреждён: ожидается непустой список"
        )

    result: dict[int, dict[str, Any]] = {}
    for index, item in enumerate(raw, start=1):
        equipment_id = item.get("id") if isinstance(item, dict) else None
        if (
            isinstance(equipment_id, bool)
            or not isinstance(equipment_id, int)
            or equipment_id <= 0
            or equipment_id in result
        ):
            raise ReportCasualtiesError(
                f"equipments.json, запись {index}: недопустимый или повторяющийся id"
            )
        result[equipment_id] = item
    return result


def _count(value: Any, label: str) -> str:
    if isinstance(value, bool) or not isinstance(value, int) or value < 0:
        raise ReportCasualtiesError(
            f"{label} должно быть целым числом не меньше нуля"
        )
    return str(value)


def load_casualty_rows(
    project_directory: Path | str,
) -> tuple[dict[str, str], ...]:
    project = Path(project_directory)
    impacts = _load_results(
        project / IMPACT_ZONES_FILE_NAME,
        "Зоны поражающих факторов не рассчитаны. "
        "Сначала выполните модуль «Зоны ПФ»",
    )
    people = _load_results(
        project / PEOPLE_FILE_NAME,
        "Число погибших и пострадавших не рассчитано. "
        "Сначала выполните модуль «Погибшие и пострадавшие»",
    )
    equipment_by_id = _load_equipment(project / "equipments.json")

    impact_by_code: dict[str, dict[str, Any]] = {}
    for index, item in enumerate(impacts, start=1):
        code = str(item.get("scenario_code", "")).strip()
        if not code or code in impact_by_code:
            raise ReportCasualtiesError(
                f"{IMPACT_ZONES_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        impact_by_code[code] = item

    rows: list[dict[str, str]] = []
    people_codes: set[str] = set()
    for index, item in enumerate(people, start=1):
        code = str(item.get("scenario_code", "")).strip()
        if not code or code in people_codes:
            raise ReportCasualtiesError(
                f"{PEOPLE_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        people_codes.add(code)
        impact = impact_by_code.get(code)
        if impact is None:
            raise ReportCasualtiesError(
                f"В {PEOPLE_FILE_NAME} найден отсутствующий "
                f"в {IMPACT_ZONES_FILE_NAME} сценарий {code}"
            )
        for field in (
            "id",
            "equipment_id",
            "equipment_name",
            "hazard_component",
            "equipment_type",
            "kind",
            "typical_scenario_line",
            "calc_code",
            "impact_status",
            "impact_values",
        ):
            if item.get(field) != impact.get(field):
                raise ReportCasualtiesError(
                    f"Результаты для сценария {code} устарели: не совпадает {field}"
                )

        equipment_id = item.get("equipment_id")
        equipment = equipment_by_id.get(equipment_id)
        if equipment is None:
            raise ReportCasualtiesError(
                f"Сценарий {code}: equipment_id={equipment_id} "
                "отсутствует в equipments.json"
            )
        for field in ("possible_dead", "possible_injured"):
            if item.get(field) != equipment.get(field):
                raise ReportCasualtiesError(
                    f"Результаты для сценария {code} устарели: не совпадает {field}"
                )

        equipment_name = str(item.get("equipment_name", "")).strip()
        component = str(item.get("hazard_component", "")).strip()
        if not equipment_name or not component:
            raise ReportCasualtiesError(
                f"Сценарий {code}: не заполнено оборудование или составляющая ОПО"
            )
        rows.append(
            {
                "code": code,
                "equipment": f"{equipment_name} ({component})",
                "fatalities": _count(
                    item.get("fatalities_count"),
                    f"Сценарий {code}: fatalities_count",
                ),
                "injured": _count(
                    item.get("injured_count"),
                    f"Сценарий {code}: injured_count",
                ),
            }
        )

    missing = set(impact_by_code) - people_codes
    if missing:
        values = ", ".join(sorted(missing))
        raise ReportCasualtiesError(
            f"В {PEOPLE_FILE_NAME} отсутствуют сценарии: {values}"
        )
    return tuple(rows)


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _cell_margins(cell: Any, value: int = 80) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _paragraph_section(document: DocumentType, paragraph_element: Any) -> Any:
    section_index = 0
    for child in document.element.body:
        if child is paragraph_element:
            break
        if child.tag == qn("w:p") and child.find(
            f"./{qn('w:pPr')}/{qn('w:sectPr')}"
        ) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def _set_table_geometry(section: Any, table: Any) -> None:
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    proportions = (0.13, 0.47, 0.20, 0.20)
    widths = [int(total_twips * value) for value in proportions[:-1]]
    widths.append(total_twips - sum(widths))
    table.autofit = False

    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row._tr.tc_lst, widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_casualties_section(
    document: DocumentType,
    rows: tuple[dict[str, str], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    section = _paragraph_section(document, marker_paragraph._p)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    headers = (
        "№ сценария",
        "Оборудование (составляющая)",
        "Количество погибших, чел.",
        "Количество пострадавших, чел.",
    )
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, bold=True, centered=True)
        _shade(cell, "D9E1F2")

    for item in rows:
        cells = table.add_row().cells
        values = (
            item["code"],
            item["equipment"],
            item["fatalities"],
            item["injured"],
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            _set_cell_text(cell, value, centered=index != 1)

    _set_repeat_header(table.rows[0])
    for row in table.rows:
        _prevent_row_split(row)
    _set_table_geometry(section, table)
    return True
