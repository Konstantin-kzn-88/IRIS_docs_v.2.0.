import json
import math
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.calculation_config import (
    CalculationConfigError,
    CalculationConfigService,
)
from iris_v2.damage_calculation import FILE_NAME as DAMAGE_FILE_NAME
from iris_v2.people_calculation import FILE_NAME as PEOPLE_FILE_NAME


MARKER = "{{DAMAGE_SECTION}}"
DAMAGE_FIELDS = (
    ("direct_losses", "Прямые потери"),
    ("liquidation_costs", "Затраты на ЛЛА"),
    ("social_losses", "Социальные потери"),
    ("indirect_damage", "Косвенный ущерб"),
    ("total_environmental_damage", "Экологический ущерб"),
    ("total_damage", "Суммарный ущерб"),
)
NOTE = (
    "ЛЛА — локализация и ликвидация аварии. "
    "Все стоимостные показатели приведены в тыс. руб."
)


class ReportDamageError(Exception):
    pass


def _load_results(path: Path, missing_message: str) -> tuple[dict[str, Any], ...]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportDamageError(missing_message) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportDamageError(f"Не удалось прочитать {path.name}") from exc
    values = raw.get("results") if isinstance(raw, dict) else None
    if not isinstance(values, list) or not values:
        raise ReportDamageError(
            f"Файл {path.name} повреждён: results должен быть непустым списком"
        )
    if any(not isinstance(item, dict) for item in values):
        raise ReportDamageError(
            f"Файл {path.name} повреждён: results содержит запись неверного формата"
        )
    return tuple(values)


def _number(value: Any, label: str) -> float:
    if (
        isinstance(value, bool)
        or not isinstance(value, (int, float))
        or not math.isfinite(float(value))
        or float(value) < 0
    ):
        raise ReportDamageError(f"{label} должно быть числом не меньше нуля")
    return float(value)


def load_damage_rows(
    project_directory: Path | str,
) -> tuple[dict[str, str], ...]:
    project = Path(project_directory)
    people = _load_results(
        project / PEOPLE_FILE_NAME,
        "Число погибших и пострадавших не рассчитано. "
        "Сначала выполните модуль «Погибшие и пострадавшие»",
    )
    damages = _load_results(
        project / DAMAGE_FILE_NAME,
        "Ущерб не рассчитан. Сначала выполните модуль «Расчёт ущерба»",
    )
    try:
        current_scale = float(CalculationConfigService().load(project)["damage_scale"])
    except (CalculationConfigError, KeyError, TypeError, ValueError) as exc:
        raise ReportDamageError("Не удалось определить текущий масштаб ущерба") from exc

    people_by_code: dict[str, dict[str, Any]] = {}
    for index, item in enumerate(people, start=1):
        code = str(item.get("scenario_code", "")).strip()
        if not code or code in people_by_code:
            raise ReportDamageError(
                f"{PEOPLE_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        people_by_code[code] = item

    rows: list[dict[str, str]] = []
    damage_codes: set[str] = set()
    for index, damage in enumerate(damages, start=1):
        code = str(damage.get("scenario_code", "")).strip()
        if not code or code in damage_codes:
            raise ReportDamageError(
                f"{DAMAGE_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        damage_codes.add(code)
        source = people_by_code.get(code)
        if source is None:
            raise ReportDamageError(
                f"В {DAMAGE_FILE_NAME} найден отсутствующий "
                f"в {PEOPLE_FILE_NAME} сценарий {code}"
            )
        for field, value in source.items():
            if damage.get(field) != value:
                raise ReportDamageError(
                    f"Результаты для сценария {code} устарели: не совпадает {field}"
                )

        stored_scale = _number(
            damage.get("damage_scale"), f"Сценарий {code}: damage_scale"
        )
        if stored_scale != current_scale:
            raise ReportDamageError(
                f"Результаты для сценария {code} устарели: изменён масштаб ущерба"
            )
        if damage.get("damage_unit") != "тыс. руб.":
            raise ReportDamageError(
                f"Сценарий {code}: неизвестная единица измерения ущерба"
            )

        equipment = str(damage.get("equipment_name", "")).strip()
        component = str(damage.get("hazard_component", "")).strip()
        if not equipment or not component:
            raise ReportDamageError(
                f"Сценарий {code}: не заполнено оборудование или составляющая ОПО"
            )
        values = {
            field: _number(damage.get(field), f"Сценарий {code}: {field}")
            for field, _ in DAMAGE_FIELDS
        }
        component_sum = sum(
            values[field] for field, _ in DAMAGE_FIELDS if field != "total_damage"
        )
        if not math.isclose(
            values["total_damage"], component_sum, rel_tol=1e-9, abs_tol=1e-6
        ):
            raise ReportDamageError(
                f"Сценарий {code}: суммарный ущерб не равен сумме составляющих"
            )

        row = {"code": code, "equipment": f"{equipment} ({component})"}
        row.update(
            {
                field: f"{values[field]:.1f}".replace(".", ",")
                for field, _ in DAMAGE_FIELDS
            }
        )
        rows.append(row)

    missing = set(people_by_code) - damage_codes
    if missing:
        values = ", ".join(sorted(missing))
        raise ReportDamageError(
            f"В {DAMAGE_FILE_NAME} отсутствуют сценарии: {values}"
        )
    return tuple(rows)


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _cell_margins(cell: Any, value: int = 60) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
    font_size: float = 8.5,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _paragraph_section(document: DocumentType, paragraph_element: Any) -> Any:
    section_index = 0
    for child in document.element.body:
        if child is paragraph_element:
            break
        if child.tag == qn("w:p") and child.find(
            f"./{qn('w:pPr')}/{qn('w:sectPr')}"
        ) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def _set_table_geometry(section: Any, table: Any) -> None:
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    proportions = (0.06, 0.23, 0.11, 0.10, 0.12, 0.12, 0.135, 0.125)
    widths = [int(total_twips * value) for value in proportions[:-1]]
    widths.append(total_twips - sum(widths))
    table.autofit = False

    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row._tr.tc_lst, widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_damage_section(
    document: DocumentType,
    rows: tuple[dict[str, str], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False

    table = document.add_table(rows=1, cols=2 + len(DAMAGE_FIELDS))
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    section = _paragraph_section(document, marker_paragraph._p)

    headers = ("№ сценария", "Оборудование (составляющая)") + tuple(
        label for _, label in DAMAGE_FIELDS
    )
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, bold=True, centered=True, font_size=7.5)
        _shade(cell, "D9E1F2")

    for item in rows:
        cells = table.add_row().cells
        values = (item["code"], item["equipment"]) + tuple(
            item[field] for field, _ in DAMAGE_FIELDS
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            _set_cell_text(cell, value, centered=index != 1)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(NOTE)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    table._tbl.addnext(note._p)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    _set_repeat_header(table.rows[0])
    for row in table.rows:
        _prevent_row_split(row)
    _set_table_geometry(section, table)
    return True
