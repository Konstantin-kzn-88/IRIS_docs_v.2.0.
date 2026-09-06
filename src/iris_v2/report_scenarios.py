import json
import math
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.calculation_cases import FILE_NAME as CASES_FILE_NAME
from iris_v2.frequency_calculation import FILE_NAME as FREQUENCY_FILE_NAME


MARKER = "{{SCENARIOS_SECTION}}"


class ReportScenariosError(Exception):
    pass


def _load_items(path: Path, field: str, missing_message: str) -> tuple[dict[str, Any], ...]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportScenariosError(missing_message) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportScenariosError(f"Не удалось прочитать {path.name}") from exc
    values = raw.get(field) if isinstance(raw, dict) else None
    if not isinstance(values, list) or not values:
        raise ReportScenariosError(
            f"Файл {path.name} повреждён: {field} должен быть непустым списком"
        )
    if any(not isinstance(item, dict) for item in values):
        raise ReportScenariosError(
            f"Файл {path.name} повреждён: {field} содержит запись неверного формата"
        )
    return tuple(values)


def load_scenario_rows(
    project_directory: Path | str,
) -> tuple[dict[str, str], ...]:
    project = Path(project_directory)
    cases = _load_items(
        project / CASES_FILE_NAME,
        "cases",
        "Расчётные сценарии не сформированы. "
        "Сначала выполните модуль «Расчётные сценарии»",
    )
    frequencies = _load_items(
        project / FREQUENCY_FILE_NAME,
        "results",
        "Частоты сценариев не рассчитаны. "
        "Сначала выполните модуль «Расчёт частот»",
    )

    case_by_code: dict[str, dict[str, Any]] = {}
    for index, case in enumerate(cases, start=1):
        code = str(case.get("scenario_code", "")).strip()
        if not code or code in case_by_code:
            raise ReportScenariosError(
                f"{CASES_FILE_NAME}, запись {index}: пустой или повторяющийся scenario_code"
            )
        case_by_code[code] = case

    rows: list[dict[str, str]] = []
    frequency_codes: set[str] = set()
    for index, result in enumerate(frequencies, start=1):
        code = str(result.get("scenario_code", "")).strip()
        if not code or code in frequency_codes:
            raise ReportScenariosError(
                f"{FREQUENCY_FILE_NAME}, запись {index}: пустой или повторяющийся scenario_code"
            )
        frequency_codes.add(code)
        case = case_by_code.get(code)
        if case is None:
            raise ReportScenariosError(
                f"В {FREQUENCY_FILE_NAME} найден отсутствующий в {CASES_FILE_NAME} сценарий {code}"
            )
        for field in ("equipment_id", "equipment_name", "hazard_component", "scenario_text"):
            if result.get(field) != case.get(field):
                raise ReportScenariosError(
                    f"Результаты для сценария {code} устарели: не совпадает {field}"
                )

        frequency = result.get("scenario_frequency")
        if (
            isinstance(frequency, bool)
            or not isinstance(frequency, (int, float))
            or not math.isfinite(float(frequency))
            or float(frequency) < 0
        ):
            raise ReportScenariosError(
                f"Сценарий {code}: scenario_frequency должно быть числом не меньше нуля"
            )
        equipment_name = str(case.get("equipment_name", "")).strip()
        component = str(case.get("hazard_component", "")).strip()
        description = str(case.get("scenario_text", "")).strip()
        if not equipment_name or not component or not description:
            raise ReportScenariosError(
                f"Сценарий {code}: не заполнено оборудование, составляющая или описание"
            )
        rows.append(
            {
                "code": code,
                "equipment": f"{equipment_name} ({component})",
                "description": description,
                "frequency": f"{float(frequency):.3E}",
            }
        )

    missing = set(case_by_code) - frequency_codes
    if missing:
        values = ", ".join(sorted(missing))
        raise ReportScenariosError(
            f"В {FREQUENCY_FILE_NAME} отсутствуют сценарии: {values}"
        )
    return tuple(rows)


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _cell_margins(cell: Any, value: int = 80) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _paragraph_section(document: DocumentType, paragraph_element: Any) -> Any:
    section_index = 0
    for child in document.element.body:
        if child is paragraph_element:
            break
        if child.tag == qn("w:p") and child.find(
            f"./{qn('w:pPr')}/{qn('w:sectPr')}"
        ) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def _set_table_geometry(section: Any, table: Any) -> None:
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    proportions = (0.09, 0.26, 0.53, 0.12)
    widths = [int(total_twips * value) for value in proportions[:-1]]
    widths.append(total_twips - sum(widths))
    table.autofit = False

    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row._tr.tc_lst, widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_scenarios_section(
    document: DocumentType,
    rows: tuple[dict[str, str], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False
    section = _paragraph_section(document, marker_paragraph._p)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    headers = (
        "№ сценария",
        "Оборудование (составляющая)",
        "Описание",
        "Частота сценария, 1/год",
    )
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, bold=True, centered=True)
        _shade(cell, "D9E1F2")
    _set_repeat_header(table.rows[0])

    for item in rows:
        cells = table.add_row().cells
        values = (
            item["code"],
            item["equipment"],
            item["description"],
            item["frequency"],
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            _set_cell_text(cell, value, centered=index in (0, 3))
    for row in table.rows:
        _prevent_row_split(row)
    _set_table_geometry(section, table)
    return True
