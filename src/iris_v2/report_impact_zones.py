import json
import math
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.hazard_factor_calculation import FILE_NAME as HAZARD_FACTOR_FILE_NAME
from iris_v2.impact_zones import FILE_NAME as IMPACT_ZONES_FILE_NAME


MARKER = "{{IMPACT_ZONES_SECTION}}"
ZONE_FIELDS = (
    ("q_10_5_m", "q=10,5"),
    ("q_7_0_m", "q=7,0"),
    ("q_4_2_m", "q=4,2"),
    ("q_1_4_m", "q=1,4"),
    ("p_28_m", "Р=28"),
    ("p_14_m", "Р=14"),
    ("p_5_m", "Р=5"),
    ("p_2_m", "Р=2"),
    ("jet_fire_length_m", "Lф"),
    ("jet_fire_diameter_m", "Dф"),
    ("lel_radius_m", "Rнкпр"),
    ("flash_fire_radius_m", "Rвсп"),
    ("lethal_radius_m", "Lпт"),
    ("threshold_radius_m", "Pпт"),
    ("dose_600_m", "Q=600"),
    ("dose_320_m", "Q=320"),
    ("dose_220_m", "Q=220"),
    ("dose_120_m", "Q=120"),
    ("spill_area_m2", "S"),
)
LEGEND = (
    "Примечание: q — интенсивность теплового излучения, кВт/м²; "
    "Р — избыточное давление взрыва ТВС, кПа; "
    "Lф — длина факела, м; Dф — диаметр факела, м; "
    "Rнкпр — радиус НКПР, м; Rвсп — радиус пожара-вспышки, м; "
    "Lпт — радиус смертельной токсодозы, м; "
    "Pпт — радиус пороговой токсодозы, м; "
    "Q — доза теплового излучения, кДж/м²; S — площадь пролива, м². "
    "Размеры зон q, Р и Q указаны в метрах."
)


class ReportImpactZonesError(Exception):
    pass


def _load_results(path: Path, missing_message: str) -> tuple[dict[str, Any], ...]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportImpactZonesError(missing_message) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportImpactZonesError(f"Не удалось прочитать {path.name}") from exc
    values = raw.get("results") if isinstance(raw, dict) else None
    if not isinstance(values, list) or not values:
        raise ReportImpactZonesError(
            f"Файл {path.name} повреждён: results должен быть непустым списком"
        )
    if any(not isinstance(item, dict) for item in values):
        raise ReportImpactZonesError(
            f"Файл {path.name} повреждён: results содержит запись неверного формата"
        )
    return tuple(values)


def _zone(value: Any, label: str) -> str:
    if value is None:
        return "—"
    if (
        isinstance(value, bool)
        or not isinstance(value, (int, float))
        or not math.isfinite(float(value))
        or float(value) < 0
    ):
        raise ReportImpactZonesError(f"{label} должно быть числом не меньше нуля")
    return f"{float(value):.1f}".replace(".", ",")


def load_impact_zone_rows(
    project_directory: Path | str,
) -> tuple[dict[str, str], ...]:
    project = Path(project_directory)
    factors = _load_results(
        project / HAZARD_FACTOR_FILE_NAME,
        "Масса поражающего фактора не рассчитана. "
        "Сначала выполните модуль «Масса ПФ»",
    )
    impacts = _load_results(
        project / IMPACT_ZONES_FILE_NAME,
        "Зоны поражающих факторов не рассчитаны. "
        "Сначала выполните модуль «Зоны ПФ»",
    )

    factor_by_code: dict[str, dict[str, Any]] = {}
    for index, item in enumerate(factors, start=1):
        code = str(item.get("scenario_code", "")).strip()
        if not code or code in factor_by_code:
            raise ReportImpactZonesError(
                f"{HAZARD_FACTOR_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        factor_by_code[code] = item

    rows: list[dict[str, str]] = []
    impact_codes: set[str] = set()
    for index, impact in enumerate(impacts, start=1):
        code = str(impact.get("scenario_code", "")).strip()
        if not code or code in impact_codes:
            raise ReportImpactZonesError(
                f"{IMPACT_ZONES_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        impact_codes.add(code)
        factor = factor_by_code.get(code)
        if factor is None:
            raise ReportImpactZonesError(
                f"В {IMPACT_ZONES_FILE_NAME} найден отсутствующий "
                f"в {HAZARD_FACTOR_FILE_NAME} сценарий {code}"
            )
        for field in ("id", "equipment_name", "hazard_component", "calc_code"):
            if impact.get(field) != factor.get(field):
                raise ReportImpactZonesError(
                    f"Результаты для сценария {code} устарели: не совпадает {field}"
                )

        equipment = str(impact.get("equipment_name", "")).strip()
        component = str(impact.get("hazard_component", "")).strip()
        if not equipment or not component:
            raise ReportImpactZonesError(
                f"Сценарий {code}: не заполнено оборудование или составляющая ОПО"
            )
        values = impact.get("impact_values")
        if not isinstance(values, dict):
            raise ReportImpactZonesError(
                f"Сценарий {code}: impact_values должен быть объектом"
            )
        spill_area = values.get("chemical_spill_area_m2")
        if spill_area is None:
            spill_area = impact.get("spill_area_m2")

        row = {"code": code, "equipment": f"{equipment} ({component})"}
        for field, _ in ZONE_FIELDS:
            value = spill_area if field == "spill_area_m2" else values.get(field)
            row[field] = _zone(value, f"Сценарий {code}: {field}")
        rows.append(row)

    missing = set(factor_by_code) - impact_codes
    if missing:
        values = ", ".join(sorted(missing))
        raise ReportImpactZonesError(
            f"В {IMPACT_ZONES_FILE_NAME} отсутствуют сценарии: {values}"
        )
    return tuple(rows)


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _cell_margins(cell: Any, value: int = 35) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(7)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _paragraph_section(document: DocumentType, paragraph_element: Any) -> Any:
    section_index = 0
    for child in document.element.body:
        if child is paragraph_element:
            break
        if child.tag == qn("w:p") and child.find(
            f"./{qn('w:pPr')}/{qn('w:sectPr')}"
        ) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def _set_table_geometry(section: Any, table: Any) -> None:
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    first_width = int(total_twips * 0.05)
    equipment_width = int(total_twips * 0.18)
    remaining = total_twips - first_width - equipment_width
    widths = [first_width, equipment_width]
    widths.extend([remaining // len(ZONE_FIELDS)] * (len(ZONE_FIELDS) - 1))
    widths.append(total_twips - sum(widths))
    table.autofit = False

    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row._tr.tc_lst, widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_impact_zones_section(
    document: DocumentType,
    rows: tuple[dict[str, str], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False

    table = document.add_table(rows=1, cols=2 + len(ZONE_FIELDS))
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    section = _paragraph_section(document, marker_paragraph._p)

    headers = ("№ сценария", "Оборудование") + tuple(
        label for _, label in ZONE_FIELDS
    )
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, bold=True, centered=True)
        _shade(cell, "D9E1F2")

    for item in rows:
        cells = table.add_row().cells
        values = (item["code"], item["equipment"]) + tuple(
            item[field] for field, _ in ZONE_FIELDS
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            _set_cell_text(cell, value, centered=index != 1)

    legend = document.add_paragraph()
    legend.paragraph_format.space_before = Pt(4)
    legend.paragraph_format.space_after = Pt(0)
    run = legend.add_run(LEGEND)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    table._tbl.addnext(legend._p)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    _set_repeat_header(table.rows[0])
    for row in table.rows:
        _prevent_row_split(row)
    _set_table_geometry(section, table)
    return True
