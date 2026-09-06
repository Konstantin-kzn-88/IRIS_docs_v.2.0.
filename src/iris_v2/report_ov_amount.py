import json
import math
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.hazard_factor_calculation import FILE_NAME as HAZARD_FACTOR_FILE_NAME
from iris_v2.release_calculation import FILE_NAME as RELEASE_FILE_NAME


MARKER = "{{OV_AMOUNT_SECTION}}"


class ReportOvAmountError(Exception):
    pass


def _load_results(path: Path, missing_message: str) -> tuple[dict[str, Any], ...]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportOvAmountError(missing_message) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportOvAmountError(f"Не удалось прочитать {path.name}") from exc
    values = raw.get("results") if isinstance(raw, dict) else None
    if not isinstance(values, list) or not values:
        raise ReportOvAmountError(
            f"Файл {path.name} повреждён: results должен быть непустым списком"
        )
    if any(not isinstance(item, dict) for item in values):
        raise ReportOvAmountError(
            f"Файл {path.name} повреждён: results содержит запись неверного формата"
        )
    return tuple(values)


def _mass(value: Any, label: str) -> float:
    if (
        isinstance(value, bool)
        or not isinstance(value, (int, float))
        or not math.isfinite(float(value))
        or float(value) < 0
    ):
        raise ReportOvAmountError(f"{label} должно быть числом не меньше нуля")
    return float(value)


def load_ov_amount_rows(
    project_directory: Path | str,
) -> tuple[dict[str, str], ...]:
    project = Path(project_directory)
    releases = _load_results(
        project / RELEASE_FILE_NAME,
        "Количество ОВ, участвующего в аварии, не рассчитано. "
        "Сначала выполните модуль «Масса в аварии»",
    )
    factors = _load_results(
        project / HAZARD_FACTOR_FILE_NAME,
        "Количество ОВ в поражающем факторе не рассчитано. "
        "Сначала выполните модуль «Масса ПФ»",
    )

    release_by_code: dict[str, dict[str, Any]] = {}
    for index, item in enumerate(releases, start=1):
        code = str(item.get("scenario_code", "")).strip()
        if not code or code in release_by_code:
            raise ReportOvAmountError(
                f"{RELEASE_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        release_by_code[code] = item

    rows: list[dict[str, str]] = []
    factor_codes: set[str] = set()
    for index, factor in enumerate(factors, start=1):
        code = str(factor.get("scenario_code", "")).strip()
        if not code or code in factor_codes:
            raise ReportOvAmountError(
                f"{HAZARD_FACTOR_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        factor_codes.add(code)
        release = release_by_code.get(code)
        if release is None:
            raise ReportOvAmountError(
                f"В {HAZARD_FACTOR_FILE_NAME} найден отсутствующий "
                f"в {RELEASE_FILE_NAME} сценарий {code}"
            )
        for field in (
            "equipment_id",
            "equipment_name",
            "hazard_component",
            "ov_in_accident_t",
        ):
            if factor.get(field) != release.get(field):
                raise ReportOvAmountError(
                    f"Результаты для сценария {code} устарели: не совпадает {field}"
                )

        equipment = str(factor.get("equipment_name", "")).strip()
        component = str(factor.get("hazard_component", "")).strip()
        if not equipment or not component:
            raise ReportOvAmountError(
                f"Сценарий {code}: не заполнено оборудование или составляющая ОПО"
            )
        accident_mass = _mass(
            factor.get("ov_in_accident_t"),
            f"Сценарий {code}: ov_in_accident_t",
        )
        factor_mass = _mass(
            factor.get("ov_in_hazard_factor_t"),
            f"Сценарий {code}: ov_in_hazard_factor_t",
        )
        rows.append(
            {
                "code": code,
                "equipment": f"{equipment} ({component})",
                "accident_mass": f"{accident_mass:.3f}".replace(".", ","),
                "factor_mass": f"{factor_mass:.3f}".replace(".", ","),
            }
        )

    missing = set(release_by_code) - factor_codes
    if missing:
        values = ", ".join(sorted(missing))
        raise ReportOvAmountError(
            f"В {HAZARD_FACTOR_FILE_NAME} отсутствуют сценарии: {values}"
        )
    return tuple(rows)


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _cell_margins(cell: Any, value: int = 80) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _paragraph_section(document: DocumentType, paragraph_element: Any) -> Any:
    section_index = 0
    for child in document.element.body:
        if child is paragraph_element:
            break
        if child.tag == qn("w:p") and child.find(
            f"./{qn('w:pPr')}/{qn('w:sectPr')}"
        ) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def _set_table_geometry(section: Any, table: Any) -> None:
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    proportions = (0.13, 0.39, 0.24, 0.24)
    widths = [int(total_twips * value) for value in proportions[:-1]]
    widths.append(total_twips - sum(widths))
    table.autofit = False

    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row._tr.tc_lst, widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_ov_amount_section(
    document: DocumentType,
    rows: tuple[dict[str, str], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    section = _paragraph_section(document, marker_paragraph._p)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    headers = (
        "№ сценария",
        "Оборудование (составляющая)",
        "Количество ОВ, участвующего в аварии, т",
        "Количество ОВ в создании поражающего фактора, т",
    )
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, bold=True, centered=True)
        _shade(cell, "D9E1F2")

    for item in rows:
        cells = table.add_row().cells
        values = (
            item["code"],
            item["equipment"],
            item["accident_mass"],
            item["factor_mass"],
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            _set_cell_text(cell, value, centered=index != 1)

    _set_repeat_header(table.rows[0])
    for row in table.rows:
        _prevent_row_split(row)
    _set_table_geometry(section, table)
    return True
