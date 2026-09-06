import json
import math
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.risk_calculation import FILE_NAME as RISK_FILE_NAME
from iris_v2.risk_summary import FILE_NAME as SUMMARY_FILE_NAME


MARKER = "{{MAX_DAMAGE_BY_COMPONENT_SECTION}}"
FIELDS = (
    ("max_direct_losses", "direct"),
    ("max_total_environmental_damage", "environmental"),
    ("max_total_damage", "total"),
)


class ReportMaxDamageError(Exception):
    pass


def _load_object(path: Path, missing_message: str) -> dict[str, Any]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportMaxDamageError(missing_message) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportMaxDamageError(f"Не удалось прочитать {path.name}") from exc
    if not isinstance(raw, dict):
        raise ReportMaxDamageError(f"Файл {path.name} повреждён: ожидается объект")
    return raw


def _number(value: Any, label: str) -> float:
    if (
        isinstance(value, bool)
        or not isinstance(value, (int, float))
        or not math.isfinite(float(value))
        or float(value) < 0
    ):
        raise ReportMaxDamageError(f"{label} должно быть числом не меньше нуля")
    return float(value)


def _same(left: float, right: float) -> bool:
    return math.isclose(left, right, rel_tol=1e-12, abs_tol=1e-9)


def load_max_damage_rows(
    project_directory: Path | str,
) -> tuple[dict[str, str], ...]:
    project = Path(project_directory)
    risk = _load_object(
        project / RISK_FILE_NAME,
        "Риски не рассчитаны. Сначала выполните модуль «Расчёт риска»",
    )
    summary = _load_object(
        project / SUMMARY_FILE_NAME,
        "Сводные показатели риска не рассчитаны. "
        "Сначала выполните модуль «Сводные показатели риска»",
    )
    results = risk.get("results")
    if not isinstance(results, list) or not results:
        raise ReportMaxDamageError(
            f"Файл {RISK_FILE_NAME} повреждён: results должен быть непустым списком"
        )

    grouped: dict[str, list[float]] = {}
    codes: set[str] = set()
    for index, item in enumerate(results, start=1):
        if not isinstance(item, dict):
            raise ReportMaxDamageError(
                f"{RISK_FILE_NAME}, запись {index}: ожидается объект"
            )
        code = str(item.get("scenario_code", "")).strip()
        component = str(item.get("hazard_component", "")).strip()
        if not code or code in codes:
            raise ReportMaxDamageError(
                f"{RISK_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        if not component:
            raise ReportMaxDamageError(
                f"Сценарий {code}: не заполнена составляющая ОПО"
            )
        if item.get("damage_unit") != "тыс. руб.":
            raise ReportMaxDamageError(
                f"Сценарий {code}: неизвестная единица измерения ущерба"
            )
        codes.add(code)
        values = (
            _number(item.get("direct_losses"), f"Сценарий {code}: direct_losses"),
            _number(
                item.get("total_environmental_damage"),
                f"Сценарий {code}: total_environmental_damage",
            ),
            _number(item.get("total_damage"), f"Сценарий {code}: total_damage"),
        )
        stored = grouped.setdefault(component, [0.0, 0.0, 0.0, 0.0])
        stored[0] += 1
        for position, value in enumerate(values, start=1):
            stored[position] = max(stored[position], value)

    if summary.get("case_count") != len(results):
        raise ReportMaxDamageError(
            f"Файл {SUMMARY_FILE_NAME} устарел: не совпадает case_count"
        )
    if summary.get("component_count") != len(grouped):
        raise ReportMaxDamageError(
            f"Файл {SUMMARY_FILE_NAME} устарел: не совпадает component_count"
        )
    if summary.get("damage_unit") != "тыс. руб.":
        raise ReportMaxDamageError(
            f"Файл {SUMMARY_FILE_NAME} содержит неизвестную единицу ущерба"
        )
    components = summary.get("components")
    if not isinstance(components, list) or len(components) != len(grouped):
        raise ReportMaxDamageError(
            f"Файл {SUMMARY_FILE_NAME} устарел: неверный перечень составляющих ОПО"
        )

    rows: list[dict[str, str]] = []
    opo_maxima = [0.0, 0.0, 0.0]
    for index, ((name, expected), component) in enumerate(
        zip(grouped.items(), components), start=1
    ):
        if not isinstance(component, dict) or component.get("hazard_component") != name:
            raise ReportMaxDamageError(
                f"Файл {SUMMARY_FILE_NAME}, составляющая {index}: неверное наименование"
            )
        if component.get("scenario_count") != int(expected[0]):
            raise ReportMaxDamageError(
                f"Файл {SUMMARY_FILE_NAME}, составляющая «{name}»: "
                "не совпадает число сценариев"
            )
        checked: list[float] = []
        for position, (field, _) in enumerate(FIELDS, start=1):
            value = _number(component.get(field), f"Составляющая «{name}»: {field}")
            if not _same(value, expected[position]):
                raise ReportMaxDamageError(
                    f"Файл {SUMMARY_FILE_NAME} устарел: максимальный ущерб "
                    f"по составляющей «{name}» не совпадает с {RISK_FILE_NAME}"
                )
            checked.append(value)
            opo_maxima[position - 1] = max(opo_maxima[position - 1], value)
        rows.append(
            {
                "component": name,
                **{
                    output: f"{value:.1f}".replace(".", ",")
                    for (_, output), value in zip(FIELDS, checked)
                },
            }
        )
    rows.append(
        {
            "component": "Максимум по ОПО",
            **{
                output: f"{value:.1f}".replace(".", ",")
                for (_, output), value in zip(FIELDS, opo_maxima)
            },
        }
    )
    return tuple(rows)


def _set_font(run: Any, size: float) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool,
    centered: bool,
    font_size: float = 9,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    _set_font(run, font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _paragraph_section(document: DocumentType, element: Any) -> Any:
    section_index = 0
    for child in document.element.body:
        if child is element:
            break
        if child.tag == qn("w:p") and child.find(
            f"./{qn('w:pPr')}/{qn('w:sectPr')}"
        ) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def _set_table_geometry(section: Any, table: Any) -> None:
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    widths = [int(total_twips * value) for value in (0.37, 0.21, 0.21)]
    widths.append(total_twips - sum(widths))
    table.autofit = False
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table._tbl.tblPr.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row._tr.tc_lst, widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_max_damage_section(
    document: DocumentType,
    rows: tuple[dict[str, str], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    headers = (
        "Составляющая ОПО",
        "Максимальные прямые потери, тыс. руб.",
        "Максимальный экологический ущерб, тыс. руб.",
        "Максимальный суммарный ущерб, тыс. руб.",
    )
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, bold=True, centered=True, font_size=8.5)
        _shade(cell, "D9E1F2")
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, item in enumerate(rows):
        cells = table.add_row().cells
        values = (
            item["component"],
            item["direct"],
            item["environmental"],
            item["total"],
        )
        is_maximum = row_index == len(rows) - 1
        for column, (cell, value) in enumerate(zip(cells, values)):
            _set_cell_text(cell, value, bold=is_maximum, centered=column > 0)
            if is_maximum:
                _shade(cell, "E2F0D9")
    marker_paragraph._element.getparent().remove(marker_paragraph._element)
    _set_table_geometry(_paragraph_section(document, table._tbl), table)
    return True
