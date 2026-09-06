import json
import math
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.risk_calculation import FILE_NAME as RISK_FILE_NAME
from iris_v2.risk_summary import FILE_NAME as SUMMARY_FILE_NAME


MARKER = "{{COLLECTIVE_RISK_SECTION}}"


class ReportCollectiveRiskError(Exception):
    pass


def _load_object(path: Path, missing_message: str) -> dict[str, Any]:
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportCollectiveRiskError(missing_message) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportCollectiveRiskError(f"Не удалось прочитать {path.name}") from exc
    if not isinstance(raw, dict):
        raise ReportCollectiveRiskError(
            f"Файл {path.name} повреждён: ожидается объект"
        )
    return raw


def _number(value: Any, label: str) -> float:
    if (
        isinstance(value, bool)
        or not isinstance(value, (int, float))
        or not math.isfinite(float(value))
        or float(value) < 0
    ):
        raise ReportCollectiveRiskError(f"{label} должно быть числом не меньше нуля")
    return float(value)


def _same(left: float, right: float) -> bool:
    return math.isclose(left, right, rel_tol=1e-12, abs_tol=0.0)


def load_collective_risk_rows(
    project_directory: Path | str,
) -> tuple[dict[str, str], ...]:
    project = Path(project_directory)
    risk = _load_object(
        project / RISK_FILE_NAME,
        "Риски не рассчитаны. Сначала выполните модуль «Расчёт риска»",
    )
    summary = _load_object(
        project / SUMMARY_FILE_NAME,
        "Сводные показатели риска не рассчитаны. "
        "Сначала выполните модуль «Сводные показатели риска»",
    )
    results = risk.get("results")
    if not isinstance(results, list) or not results:
        raise ReportCollectiveRiskError(
            f"Файл {RISK_FILE_NAME} повреждён: results должен быть непустым списком"
        )

    grouped: dict[str, list[float]] = {}
    codes: set[str] = set()
    for index, item in enumerate(results, start=1):
        if not isinstance(item, dict):
            raise ReportCollectiveRiskError(
                f"{RISK_FILE_NAME}, запись {index}: ожидается объект"
            )
        code = str(item.get("scenario_code", "")).strip()
        component = str(item.get("hazard_component", "")).strip()
        if not code or code in codes:
            raise ReportCollectiveRiskError(
                f"{RISK_FILE_NAME}, запись {index}: "
                "пустой или повторяющийся scenario_code"
            )
        if not component:
            raise ReportCollectiveRiskError(f"Сценарий {code}: не заполнена составляющая ОПО")
        codes.add(code)
        values = grouped.setdefault(component, [0.0, 0.0, 0.0])
        values[0] += 1
        values[1] += _number(
            item.get("collective_risk_fatalities"),
            f"Сценарий {code}: collective_risk_fatalities",
        )
        values[2] += _number(
            item.get("collective_risk_injured"),
            f"Сценарий {code}: collective_risk_injured",
        )

    if summary.get("case_count") != len(results):
        raise ReportCollectiveRiskError(
            f"Файл {SUMMARY_FILE_NAME} устарел: не совпадает case_count"
        )
    if summary.get("component_count") != len(grouped):
        raise ReportCollectiveRiskError(
            f"Файл {SUMMARY_FILE_NAME} устарел: не совпадает component_count"
        )
    if summary.get("risk_unit") != "1/год":
        raise ReportCollectiveRiskError(
            f"Файл {SUMMARY_FILE_NAME} содержит неизвестную единицу риска"
        )
    components = summary.get("components")
    if not isinstance(components, list) or len(components) != len(grouped):
        raise ReportCollectiveRiskError(
            f"Файл {SUMMARY_FILE_NAME} устарел: неверный перечень составляющих ОПО"
        )

    rows: list[dict[str, str]] = []
    for index, ((name, expected), stored) in enumerate(
        zip(grouped.items(), components), start=1
    ):
        if not isinstance(stored, dict) or stored.get("hazard_component") != name:
            raise ReportCollectiveRiskError(
                f"Файл {SUMMARY_FILE_NAME}, составляющая {index}: неверное наименование"
            )
        if stored.get("scenario_count") != int(expected[0]):
            raise ReportCollectiveRiskError(
                f"Файл {SUMMARY_FILE_NAME}, составляющая «{name}»: "
                "не совпадает число сценариев"
            )
        fatalities = _number(
            stored.get("collective_risk_fatalities"),
            f"Составляющая «{name}»: collective_risk_fatalities",
        )
        injured = _number(
            stored.get("collective_risk_injured"),
            f"Составляющая «{name}»: collective_risk_injured",
        )
        if not _same(fatalities, expected[1]) or not _same(injured, expected[2]):
            raise ReportCollectiveRiskError(
                f"Файл {SUMMARY_FILE_NAME} устарел: риск по составляющей «{name}» "
                f"не совпадает с {RISK_FILE_NAME}"
            )
        rows.append(
            {
                "component": name,
                "fatalities": f"{fatalities:.3E}",
                "injured": f"{injured:.3E}",
            }
        )

    total_fatalities = sum(values[1] for values in grouped.values())
    total_injured = sum(values[2] for values in grouped.values())
    stored_total_fatalities = _number(
        summary.get("total_collective_risk_fatalities"),
        "total_collective_risk_fatalities",
    )
    stored_total_injured = _number(
        summary.get("total_collective_risk_injured"),
        "total_collective_risk_injured",
    )
    if not _same(stored_total_fatalities, total_fatalities) or not _same(
        stored_total_injured, total_injured
    ):
        raise ReportCollectiveRiskError(
            f"Файл {SUMMARY_FILE_NAME} устарел: итоговый коллективный риск "
            f"не совпадает с {RISK_FILE_NAME}"
        )
    rows.append(
        {
            "component": "Итого по ОПО",
            "fatalities": f"{stored_total_fatalities:.3E}",
            "injured": f"{stored_total_injured:.3E}",
        }
    )
    return tuple(rows)


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_text(cell: Any, value: str, *, bold: bool, centered: bool) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _paragraph_section(document: DocumentType, paragraph_element: Any) -> Any:
    section_index = 0
    for child in document.element.body:
        if child is paragraph_element:
            break
        if child.tag == qn("w:p") and child.find(
            f"./{qn('w:pPr')}/{qn('w:sectPr')}"
        ) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def _set_table_geometry(section: Any, table: Any) -> None:
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    widths = (int(total_twips * 0.46), int(total_twips * 0.27))
    widths += (total_twips - sum(widths),)
    table.autofit = False
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table._tbl.tblPr.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row._tr.tc_lst, widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_collective_risk_section(
    document: DocumentType,
    rows: tuple[dict[str, str], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    headers = (
        "Составляющая ОПО",
        "Коллективный риск гибели, чел./год",
        "Коллективный риск травмирования, чел./год",
    )
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, bold=True, centered=True)
        _shade(cell, "D9E1F2")
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, item in enumerate(rows):
        cells = table.add_row().cells
        values = (item["component"], item["fatalities"], item["injured"])
        is_total = row_index == len(rows) - 1
        for column, (cell, value) in enumerate(zip(cells, values)):
            _set_cell_text(cell, value, bold=is_total, centered=column > 0)
            if is_total:
                _shade(cell, "E2F0D9")
    marker_paragraph._element.getparent().remove(marker_paragraph._element)
    _set_table_geometry(_paragraph_section(document, table._tbl), table)
    return True
