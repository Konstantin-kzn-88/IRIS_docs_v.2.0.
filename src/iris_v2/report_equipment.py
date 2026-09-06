import json
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from iris_v2.equipment import JSON_FILE_NAME, PIPELINE_TYPES


MARKER = "{{EQUIPMENT_SECTION}}"


class ReportEquipmentError(Exception):
    pass


def load_project_equipment(
    project_directory: Path | str,
) -> tuple[dict[str, Any], ...]:
    path = Path(project_directory) / JSON_FILE_NAME
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ReportEquipmentError(
            "Оборудование проекта не импортировано. "
            "Сначала заполните модуль «Оборудование»"
        ) from exc
    except (OSError, json.JSONDecodeError) as exc:
        raise ReportEquipmentError(
            f"Не удалось прочитать {JSON_FILE_NAME}"
        ) from exc
    if not isinstance(raw, list) or not raw:
        raise ReportEquipmentError(
            "В equipments.json нет оборудования. "
            "Сначала заполните модуль «Оборудование»"
        )
    if any(not isinstance(item, dict) for item in raw):
        raise ReportEquipmentError(
            "Файл equipments.json повреждён: ожидается список объектов"
        )
    return tuple(raw)


def _is_empty(value: Any) -> bool:
    return value is None or (isinstance(value, str) and value.strip() in ("", "-"))


def _number(value: int | float) -> str:
    return format(value, ".12g").replace(".", ",")


def _value(value: Any, unit: str = "") -> str:
    if isinstance(value, bool):
        text = "да" if value else "нет"
    elif isinstance(value, (int, float)):
        text = _number(value)
    else:
        text = str(value).strip()
    return f"{text} {unit}".strip()


def _append(
    rows: list[tuple[str, str]],
    item: dict[str, Any],
    key: str,
    label: str,
    unit: str = "",
) -> None:
    value = item.get(key)
    if not _is_empty(value):
        rows.append((label, _value(value, unit)))


def equipment_sections(
    item: dict[str, Any],
    substance_names: dict[int, str],
) -> list[tuple[str, list[tuple[str, str]]]]:
    general: list[tuple[str, str]] = []
    _append(general, item, "equipment_name", "Наименование оборудования")
    _append(general, item, "hazard_component", "Составляющая ОПО")

    substance_id = item.get("substance_id")
    if (
        isinstance(substance_id, bool)
        or not isinstance(substance_id, int)
        or substance_id not in substance_names
    ):
        raise ReportEquipmentError(
            f"Оборудование {item.get('id', '—')}: substance_id отсутствует "
            "в substances.json"
        )
    general.append(("Опасное вещество", substance_names[substance_id]))
    _append(general, item, "phase_state", "Фазовое состояние вещества")

    geometry: list[tuple[str, str]] = []
    equipment_type = item.get("equipment_type")
    if equipment_type in PIPELINE_TYPES:
        _append(
            geometry,
            item,
            "total_length_m",
            "Общая протяжённость для расчёта частоты",
            "м",
        )
        _append(
            geometry,
            item,
            "accident_section_length_m",
            "Длина аварийного участка",
            "м",
        )
        _append(geometry, item, "diameter_mm", "Наружный диаметр", "мм")
        _append(geometry, item, "wall_thickness_mm", "Толщина стенки", "мм")
    else:
        _append(
            geometry,
            item,
            "equipment_count",
            "Количество оборудования для расчёта частоты",
            "шт.",
        )
        _append(geometry, item, "volume_m3", "Объём оборудования", "м³")
        _append(geometry, item, "fill_fraction", "Степень заполнения")

    regime: list[tuple[str, str]] = []
    _append(regime, item, "pressure_mpa", "Давление", "МПа")
    _append(
        regime,
        item,
        "substance_temperature_c",
        "Температура вещества",
        "°C",
    )

    spill: list[tuple[str, str]] = []
    _append(
        spill,
        item,
        "spill_coefficient",
        "Коэффициент растекания",
        "м⁻¹",
    )
    _append(spill, item, "spill_area_m2", "Заданная площадь разлива", "м²")
    _append(spill, item, "shutdown_time_s", "Время отключения оборудования", "с")
    _append(spill, item, "evaporation_time_s", "Время испарения из пролива", "с")

    sections = [
        ("Общие сведения", general),
        ("Количество и геометрия", geometry),
        ("Режим", regime),
        ("Разлив и испарение", spill),
    ]
    return [(title, rows) for title, rows in sections if rows]


def _set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _cell_margins(cell: Any, value: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: Any,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _set_table_geometry(document: DocumentType, table: Any) -> None:
    section = document.sections[0]
    total_twips = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    indent_twips = 100
    total_twips -= indent_twips
    left_width = int(total_twips * 0.42)
    widths = (left_width, total_twips - left_width)
    table.autofit = False

    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(indent_twips))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        cells = row._tr.tc_lst
        row_widths = (total_twips,) if len(cells) == 1 else widths
        for cell, width in zip(cells, row_widths):
            cell_width = cell.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def render_equipment_section(
    document: DocumentType,
    equipment: tuple[dict[str, Any], ...],
    substances: tuple[dict[str, Any], ...],
) -> bool:
    marker_paragraph = next(
        (paragraph for paragraph in document.paragraphs if MARKER in paragraph.text),
        None,
    )
    if marker_paragraph is None:
        return False
    if not equipment:
        raise ReportEquipmentError(
            "В equipments.json нет оборудования. "
            "Сначала заполните модуль «Оборудование»"
        )

    substance_names = {
        item["id"]: str(item.get("name", "")).strip()
        for item in substances
        if isinstance(item.get("id"), int) and str(item.get("name", "")).strip()
    }

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    marker_paragraph._p.addnext(table._tbl)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    header = table.rows[0].cells
    _set_cell_text(header[0], "Параметр", bold=True, centered=True)
    _set_cell_text(header[1], "Значение", bold=True, centered=True)
    _shade(header[0], "D9E1F2")
    _shade(header[1], "D9E1F2")
    _set_repeat_header(table.rows[0])

    for item in equipment:
        title = str(item.get("equipment_name", "")).strip() or "—"
        title_cells = table.add_row().cells
        title_cell = title_cells[0].merge(title_cells[1])
        _set_cell_text(title_cell, title, bold=True)
        _shade(title_cell, "EDEDED")
        for section_title, rows in equipment_sections(item, substance_names):
            section_cells = table.add_row().cells
            section_cell = section_cells[0].merge(section_cells[1])
            _set_cell_text(section_cell, section_title, bold=True)
            _shade(section_cell, "F5F5F5")
            for label, value in rows:
                cells = table.add_row().cells
                _set_cell_text(cells[0], label)
                _set_cell_text(cells[1], value)

    for row in table.rows:
        _prevent_row_split(row)
    _set_table_geometry(document, table)
    return True
